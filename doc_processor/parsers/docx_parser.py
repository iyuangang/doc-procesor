"""
docx文档解析器模块
"""

import os
import re
import shutil
import tempfile
from functools import lru_cache
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table as DocxTable
from lxml import etree

from doc_processor.exceptions import DocumentError
from doc_processor.models import DocumentNode, DocumentStructure
from doc_processor.utils import (
    TimingContext,
    clean_text,
    create_logger,
    extract_batch_number,
    extract_count_from_text,
)

# 创建日志记录器
logger = create_logger(__name__)


class DocxParser:
    """
    docx文档解析器

    Attributes:
        doc_path: 文档路径
        doc: 文档对象
        doc_structure: 文档结构
        batch_number: 批次号
        declared_count: 声明的总记录数
    """

    def __init__(self, doc_path: str) -> None:
        """
        初始化解析器

        Args:
            doc_path: 文档路径
        """
        self.doc_path = doc_path
        self.doc_structure = DocumentStructure()
        self.batch_number: Optional[str] = None
        self.declared_count: Optional[int] = None
        self._load_document()

    def _load_document(self) -> None:
        """
        安全加载文档，处理大文件

        Raises:
            DocumentError: 加载文档失败
        """
        try:
            file_size = os.path.getsize(self.doc_path)
            logger.info(f"加载文档 {self.doc_path}, 大小: {file_size/1024/1024:.2f}MB")

            if file_size > 100 * 1024 * 1024:  # 100MB
                logger.warning(f"文档大小超过100MB，使用临时文件处理")
                with tempfile.NamedTemporaryFile(delete=False) as tmp:
                    shutil.copy2(self.doc_path, tmp.name)
                    self.doc = Document(tmp.name)
                    os.unlink(tmp.name)
            else:
                self.doc = Document(self.doc_path)
        except Exception as e:
            logger.error(f"加载文档失败: {str(e)}")
            raise DocumentError(f"无法加载文档: {str(e)}", self.doc_path)

    def parse(self) -> DocumentStructure:
        """
        解析文档

        Returns:
            文档结构
        """
        with TimingContext("解析文档"):
            # 预处理 - 提取批次号和总记录数
            self._extract_batch_info()

            # 处理文档中的所有元素
            for element in self.doc.element.body:
                try:
                    # 处理段落
                    if element.tag.endswith("p"):
                        self._process_paragraph(element)
                    # 处理表格
                    elif element.tag.endswith("tbl"):
                        self._process_table(element)
                except Exception as e:
                    logger.error(f"处理元素出错: {str(e)}")
                    continue

        return self.doc_structure

    def _extract_batch_info(self) -> None:
        """提取批次号和总记录数"""
        # 从前几个段落提取批次号和总记录数
        search_limit = min(20, len(self.doc.paragraphs))

        for i, para in enumerate(self.doc.paragraphs[:search_limit]):
            text = para.text.strip()
            if not text:
                continue

            # 提取批次号
            if not self.batch_number and "批" in text:
                self.batch_number = extract_batch_number(text)
                if self.batch_number:
                    self.doc_structure.set_batch_number(self.batch_number)
                    logger.info(f"提取到批次号: {self.batch_number}")
                    self.doc_structure.add_node(
                        f"第{self.batch_number}批", "batch", level=0
                    )

            # 提取总记录数
            if not self.declared_count and any(
                word in text for word in ["共计", "总计", "合计"]
            ):
                self.declared_count = extract_count_from_text(text)
                if self.declared_count:
                    logger.info(f"提取到总记录数: {self.declared_count}")

    def _process_paragraph(self, element: Any) -> None:
        """
        处理段落元素

        Args:
            element: 段落元素
        """
        # 获取段落文本
        text = ""
        for t in element.xpath(".//w:t"):
            if t.text:
                text += t.text

        text = text.strip()
        if not text:
            return

        # 提取批次号
        if not self.batch_number and "批" in text:
            self.batch_number = extract_batch_number(text)
            if self.batch_number:
                self.doc_structure.set_batch_number(self.batch_number)
                logger.info(f"提取到批次号: {self.batch_number}")
                self.doc_structure.add_node(
                    f"第{self.batch_number}批", "batch", level=0
                )
                return

        # 处理文档结构
        if "一、节能型汽车" in text:
            self.doc_structure.add_node("节能型汽车", "section", content=text)
        elif "二、新能源汽车" in text:
            self.doc_structure.add_node("新能源汽车", "section", content=text)
        elif text.startswith("（") and "）" in text:
            self.doc_structure.add_node(text.strip(), "subsection", content=text)
        elif text.startswith(("1.", "2.", "3.", "4.", "5.")):
            self.doc_structure.add_node(text.strip(), "numbered_section", content=text)
        elif text.startswith("（") and any(num in text for num in "123456789"):
            self.doc_structure.add_node(
                text.strip(), "numbered_subsection", content=text
            )
        elif "勘误" in text or "说明" in text:
            self.doc_structure.add_node(
                text[:20] + "..." if len(text) > 20 else text, "note", content=text
            )
        elif "更正" in text or "修改" in text:
            self.doc_structure.add_node(
                text[:20] + "..." if len(text) > 20 else text,
                "correction",
                content=text,
            )
        else:
            self.doc_structure.add_node(
                text[:20] + "..." if len(text) > 20 else text, "text", content=text
            )

    def _process_table(self, element: Any) -> None:
        """
        处理表格元素

        Args:
            element: 表格元素
        """
        # 查找对应的表格对象
        for i, table in enumerate(self.doc.tables):
            if table._element is element:
                if not table.rows:
                    return

                # 使用索引作为表格ID
                table_id = i + 1

                # 提取表格基本信息
                rows_count = len(table.rows)
                cols_count = len(table.rows[0].cells) if rows_count > 0 else 0

                # 添加表格节点
                self.doc_structure.add_node(
                    f"表格 {table_id}",
                    "table",
                    metadata={
                        "rows": rows_count,
                        "columns": cols_count,
                        "id": table_id,
                    },
                )
                break

    def get_paragraphs(self) -> List[str]:
        """
        获取文档中的段落文本

        Returns:
            段落文本列表
        """
        return [para.text.strip() for para in self.doc.paragraphs if para.text.strip()]

    def get_tables(self) -> List[DocxTable]:
        """
        获取文档中的表格

        Returns:
            表格列表
        """
        return self.doc.tables

    def get_batch_number(self) -> Optional[str]:
        """
        获取批次号

        Returns:
            批次号，如果未找到则返回None
        """
        return self.batch_number

    def get_declared_count(self) -> Optional[int]:
        """
        获取声明的总记录数

        Returns:
            总记录数，如果未找到则返回None
        """
        return self.declared_count

    def extract_doc_content(self) -> Tuple[List[str], List[Dict[str, str]]]:
        """
        提取文档中除表格外的内容，并分离额外信息

        Returns:
            段落文本列表和额外信息列表的元组
        """
        paragraphs: List[str] = []
        extra_info: List[Dict[str, str]] = []
        current_section: Optional[str] = None
        batch_found = False
        batch_number = None

        # 额外信息的标识词和对应类型
        info_types: Dict[str, str] = {
            "勘误": "勘误",
            "关于": "政策",
            "符合": "说明",
            "技术要求": "说明",
            "自动转入": "说明",
            "第二部分": "说明",
        }

        # 用于收集连续的额外信息文本
        current_extra_info: Optional[Dict[str, str]] = None

        def save_current_extra_info() -> None:
            """保存当前的额外信息"""
            nonlocal current_extra_info
            if current_extra_info:
                # 清理和规范化内容
                content = current_extra_info["content"]
                # 移除多余的空白字符
                content = re.sub(r"\s+", " ", content)
                # 移除换行符
                content = content.replace("\n", " ")
                current_extra_info["content"] = content.strip()

                # 添加批次号
                if batch_number:
                    current_extra_info["batch"] = batch_number

                # 检查是否需要合并相同类型和章节的信息
                for info in extra_info:
                    if (
                        info["type"] == current_extra_info["type"]
                        and info["section"] == current_extra_info["section"]
                    ):
                        info["content"] = (
                            info["content"] + " " + current_extra_info["content"]
                        )
                        current_extra_info = None
                        return

                extra_info.append(current_extra_info)
                current_extra_info = None

        # 遍历文档段落
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                # 如果遇到空行，保存当前的额外信息
                if current_extra_info:
                    save_current_extra_info()
                continue

            # 检查批次号
            if not batch_found and "批" in text:
                extracted_batch = extract_batch_number(text)
                if extracted_batch:
                    batch_number = extracted_batch
                    paragraphs.append(text)  # 将批次号信息放在最前面
                    batch_found = True
                    continue

            # 识别主要分类
            if text.startswith("一、") or text.startswith("二、"):
                save_current_extra_info()
                current_section = text
                paragraphs.append(text)
            # 识别子分类
            elif text.startswith("（"):
                save_current_extra_info()
                current_section = text
                paragraphs.append(text)
            # 识别额外信息
            elif any(marker in text for marker in info_types.keys()):
                # 如果当前文本包含新的标识词，保存之前的信息并创建新的
                if current_extra_info:
                    save_current_extra_info()

                # 创建新的额外信息
                info_type = next(
                    (t for m, t in info_types.items() if m in text), "其他"
                )
                current_extra_info = {
                    "section": current_section or "文档说明",
                    "type": info_type,
                    "content": text,
                }
            # 如果当前有未处理的额外信息，将文本追加到内容中
            elif current_extra_info is not None:
                current_extra_info["content"] = (
                    current_extra_info["content"] + " " + text
                )
            else:
                paragraphs.append(text)

        # 保存最后一条未处理的额外信息
        save_current_extra_info()

        return paragraphs, extra_info
