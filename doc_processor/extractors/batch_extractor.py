"""
批次号提取器模块
"""

import re
from typing import List, Optional, Tuple

from docx.document import Document as DocxDocument

from doc_processor.exceptions import ExtractionError
from doc_processor.extractors.base import BaseExtractor
from doc_processor.utils import extract_batch_number, extract_count_from_text


class BatchExtractor(BaseExtractor[Tuple[Optional[str], Optional[int]]]):
    """
    批次号提取器

    Attributes:
        doc: 文档对象
        max_paragraphs: 最大搜索段落数
    """

    def __init__(
        self, doc: DocxDocument, max_paragraphs: int = 20, verbose: bool = False
    ) -> None:
        """
        初始化批次号提取器

        Args:
            doc: 文档对象
            max_paragraphs: 最大搜索段落数
            verbose: 是否显示详细信息
        """
        super().__init__("批次号提取器", verbose)
        self.doc = doc
        self.max_paragraphs = max_paragraphs

    def extract(self) -> Tuple[Optional[str], Optional[int]]:
        """
        提取批次号和总记录数

        Returns:
            批次号和总记录数的元组

        Raises:
            ExtractionError: 提取失败
        """
        try:
            batch_number = None
            declared_count = None

            # 从前几个段落提取批次号和总记录数
            search_limit = min(self.max_paragraphs, len(self.doc.paragraphs))

            for i, para in enumerate(self.doc.paragraphs[:search_limit]):
                text = para.text.strip()
                if not text:
                    continue

                # 提取批次号
                if not batch_number and "批" in text:
                    batch_number = extract_batch_number(text)
                    if batch_number:
                        self.log_info(f"提取到批次号: {batch_number}")

                # 提取总记录数
                if not declared_count and any(
                    word in text for word in ["共计", "总计", "合计"]
                ):
                    declared_count = extract_count_from_text(text)
                    if declared_count:
                        self.log_info(f"提取到总记录数: {declared_count}")

                # 如果都找到了，可以提前结束
                if batch_number and declared_count:
                    break

            # 记录结果
            if not batch_number:
                self.log_warning("未提取到批次号")
            if not declared_count:
                self.log_warning("未提取到总记录数")

            return batch_number, declared_count

        except Exception as e:
            raise ExtractionError(f"提取批次号和总记录数失败: {str(e)}")


class ContentExtractor(BaseExtractor[Tuple[List[str], List[dict]]]):
    """
    内容提取器，提取文档中的段落和额外信息

    Attributes:
        doc: 文档对象
    """

    def __init__(self, doc: DocxDocument, verbose: bool = False) -> None:
        """
        初始化内容提取器

        Args:
            doc: 文档对象
            verbose: 是否显示详细信息
        """
        super().__init__("内容提取器", verbose)
        self.doc = doc

    def extract(self) -> Tuple[List[str], List[dict]]:
        """
        提取文档中的段落和额外信息

        Returns:
            段落和额外信息的元组

        Raises:
            ExtractionError: 提取失败
        """
        try:
            paragraphs: List[str] = []
            extra_info: List[dict] = []
            current_section: Optional[str] = None
            batch_found = False
            batch_number = None

            # 额外信息的标识词和对应类型
            info_types = {
                "勘误": "勘误",
                "关于": "政策",
                "符合": "说明",
                "技术要求": "说明",
                "自动转入": "说明",
                "第二部分": "说明",
            }

            # 用于收集连续的额外信息文本
            current_extra_info: Optional[dict] = None

            def save_current_extra_info() -> None:
                """保存当前的额外信息"""
                nonlocal current_extra_info, extra_info
                if current_extra_info:
                    # 清理和规范化内容
                    content = current_extra_info["content"]
                    # 移除多余的空白字符
                    content = re.sub(r"\s+", " ", content)
                    # 移除换行符
                    content = content.replace("\n", " ")
                    current_extra_info["content"] = content.strip()

                    # 添加批次号
                    if batch_number:
                        current_extra_info["batch"] = batch_number

                    # 检查是否需要合并相同类型和章节的信息
                    for info in extra_info:
                        if (
                            info["type"] == current_extra_info["type"]
                            and info["section"] == current_extra_info["section"]
                        ):
                            info["content"] = (
                                info["content"] + " " + current_extra_info["content"]
                            )
                            current_extra_info = None
                            return

                    extra_info.append(current_extra_info)
                    current_extra_info = None

            # 遍历文档段落
            for para in self.doc.paragraphs:
                text = para.text.strip()
                if not text:
                    # 如果遇到空行，保存当前的额外信息
                    if current_extra_info:
                        save_current_extra_info()
                    continue

                # 检查批次号
                if not batch_found and "批" in text:
                    extracted_batch = extract_batch_number(text)
                    if extracted_batch:
                        batch_number = extracted_batch
                        paragraphs.append(text)  # 将批次号信息放在最前面
                        batch_found = True
                        continue

                # 识别主要分类
                if text.startswith("一、") or text.startswith("二、"):
                    save_current_extra_info()
                    current_section = text
                    paragraphs.append(text)
                # 识别子分类
                elif text.startswith("（"):
                    save_current_extra_info()
                    current_section = text
                    paragraphs.append(text)
                # 识别额外信息
                elif any(marker in text for marker in info_types.keys()):
                    # 如果当前文本包含新的标识词，保存之前的信息并创建新的
                    if current_extra_info:
                        save_current_extra_info()

                    # 创建新的额外信息
                    info_type = next(
                        (t for m, t in info_types.items() if m in text), "其他"
                    )
                    current_extra_info = {
                        "section": current_section or "文档说明",
                        "type": info_type,
                        "content": text,
                    }
                # 如果当前有未处理的额外信息，将文本追加到内容中
                elif current_extra_info is not None:
                    current_extra_info["content"] = (
                        current_extra_info["content"] + " " + text
                    )
                else:
                    paragraphs.append(text)

            # 保存最后一条未处理的额外信息
            save_current_extra_info()

            self.log_info(
                f"提取到 {len(paragraphs)} 个段落和 {len(extra_info)} 条额外信息"
            )
            return paragraphs, extra_info

        except Exception as e:
            raise ExtractionError(f"提取文档内容失败: {str(e)}")
