"""
测试工具函数 - 提供通用测试辅助功能
"""

import os
import random
import string
import io
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_sample_document(
    output_path: str,
    batch_number: str = "1",
    car_records: Optional[List[Dict[str, Any]]] = None,
) -> str:
    """
    创建用于测试的示例Word文档

    Args:
        output_path: 输出文件路径
        batch_number: 批次号
        car_records: 车辆记录列表，如果为None则创建默认记录

    Returns:
        创建的文档路径
    """
    # 创建默认记录
    if car_records is None:
        car_records = [
            {
                "企业名称": "测试企业A",
                "品牌": "测试品牌X",
                "型号": "测试型号1",
                "排量": "1.5L",
                "变速器类型": "AT",
                "档位数": "6",
            },
            {
                "企业名称": "测试企业B",
                "品牌": "测试品牌Y",
                "型号": "测试型号2",
                "排量": "2.0L",
                "变速器类型": "CVT",
            },
            {
                "企业名称": "测试企业C",
                "品牌": "测试品牌Z",
                "型号": "测试型号3",
                "排量": "1.8L",
                "变速器类型": "MT",
                "档位数": "5",
            },
        ]

    # 创建文档
    doc = Document()

    # 设置文档样式
    style = doc.styles["Normal"]
    font = style.font
    font.name = "宋体"
    font.size = Pt(12)

    # 添加标题
    title = doc.add_heading(f"第{batch_number}批车辆公告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加正文
    doc.add_paragraph(f"第{batch_number}批车辆公告包含以下车型:")

    # 添加分类段落
    p1 = doc.add_paragraph()
    p1.add_run("节能型汽车").bold = True

    p2 = doc.add_paragraph()
    p2.add_run("（一）轿车").bold = True

    # 添加说明
    doc.add_paragraph("以下为本批次节能型轿车，共计{0}款。".format(len(car_records)))

    # 添加表格
    table = doc.add_table(rows=1, cols=len(car_records[0].keys()) + 1)
    table.style = "Table Grid"

    # 设置表头
    header_cells = table.rows[0].cells
    header_cells[0].text = "序号"
    for i, key in enumerate(car_records[0].keys()):
        header_cells[i + 1].text = key

    # 添加数据行
    for i, record in enumerate(car_records):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        for j, key in enumerate(record.keys()):
            row[j + 1].text = str(record.get(key, ""))

    # 保存文档
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)

    return output_path


def generate_random_string(length: int = 10) -> str:
    """生成随机字符串"""
    return "".join(random.choices(string.ascii_letters + string.digits, k=length))


def create_temp_file(
    directory: str,
    filename: Optional[str] = None,
    content: str = "",
    extension: str = ".txt",
) -> str:
    """
    在指定目录创建临时文件

    Args:
        directory: 目录路径
        filename: 文件名，如果为None则生成随机文件名
        content: 文件内容
        extension: 文件扩展名

    Returns:
        创建的文件路径
    """
    os.makedirs(directory, exist_ok=True)

    if filename is None:
        filename = generate_random_string() + extension
    elif not filename.endswith(extension):
        filename += extension

    filepath = os.path.join(directory, filename)

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)

    return filepath


def simulate_docx_file_content() -> bytes:
    """
    模拟docx文件内容，用于测试

    Returns:
        文件内容字节
    """
    # 创建Word文档
    doc = Document()
    doc.add_heading("测试文档", 0)
    doc.add_paragraph("这是一个用于测试的文档")

    # 将文档保存到内存中
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    content = f.read()

    return content
