"""
控制台输出模块 - 提供控制台显示和格式化功能
"""

import textwrap
import time
from typing import Dict, Any, List, Set, Optional, Tuple

from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.progress import (
    Progress,
    SpinnerColumn,
    BarColumn,
    TimeRemainingColumn,
    TimeElapsedColumn,
)
from rich.text import Text
from rich.tree import Tree
from rich.layout import Layout
from rich.columns import Columns
from rich import box

from ..models.document_node import DocumentNode, DocumentStructure

# 创建控制台对象
console = Console()


def display_statistics(
    total_count: int, energy_saving_count: int, new_energy_count: int, output_file: str
) -> None:
    """
    显示处理统计信息

    Args:
        total_count: 总记录数
        energy_saving_count: 节能型汽车数量
        new_energy_count: 新能源汽车数量
        output_file: 输出文件路径
    """
    # 在显示表格前添加标题，表明这是关键信息
    console.print()
    console.print("[bold cyan]📊 关键信息：处理统计报告[/bold cyan]")

    # 创建统计表格
    stats_table = Table(
        title="📊 处理统计报告",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    stats_table.add_column("项目", style="cyan")
    stats_table.add_column("数值", justify="right", style="green")
    stats_table.add_column("占比", justify="right", style="yellow")

    # 计算百分比
    energy_saving_percent = (
        energy_saving_count / total_count * 100 if total_count > 0 else 0
    )
    new_energy_percent = new_energy_count / total_count * 100 if total_count > 0 else 0

    # 添加行
    stats_table.add_row("📝 总记录数", f"{total_count:,}", "100%")
    stats_table.add_row(
        "🚗 节能型汽车", f"{energy_saving_count:,}", f"{energy_saving_percent:.1f}%"
    )
    stats_table.add_row(
        "⚡ 新能源汽车", f"{new_energy_count:,}", f"{new_energy_percent:.1f}%"
    )
    stats_table.add_row("💾 输出文件", output_file, "")

    # 显示表格
    console.print(stats_table)
    console.print()


def display_doc_content(doc_structure: DocumentStructure) -> None:
    """
    使用树形结构显示文档内容

    Args:
        doc_structure: 文档结构对象
    """

    def get_node_style(node: DocumentNode) -> Tuple[str, str]:
        """获取节点的样式和图标"""
        styles = {
            "root": ("bold blue", "📑"),
            "section": ("bold cyan", "📌"),
            "subsection": ("bold yellow", "📎"),
            "numbered_section": ("bold green", "🔢"),
            "numbered_subsection": ("bold magenta", "📍"),
            "table": ("bold blue", "📊"),
            "text": ("white", "📝"),
            "note": ("bold magenta", "ℹ️"),
            "correction": ("bold red", "⚠️"),
        }
        return styles.get(node.node_type, ("white", "•"))

    def add_node_to_tree(tree: Tree, node: DocumentNode) -> None:
        """递归添加节点到树中"""
        style, icon = get_node_style(node)

        # 构建节点标题
        title = f"{icon} {node.title}"
        if node.batch_number and node.level <= 1:
            title += f" [dim](第{node.batch_number}批)[/dim]"

        # 创建节点
        branch = tree.add(f"[{style}]{title}[/{style}]")

        # 添加内容（如果有且与标题不同）
        if node.content and node.content != node.title:
            content_lines = textwrap.wrap(node.content, width=100)
            for line in content_lines:
                branch.add(f"[dim]{line}[/dim]")

        # 添加元数据（如果有）
        if node.metadata:
            meta_branch = branch.add("[dim]元数据[/dim]")
            for key, value in node.metadata.items():
                meta_branch.add(f"[dim]{key}: {value}[/dim]")

        # 递归处理子节点
        for child in node.children:
            add_node_to_tree(branch, child)

    # 创建主树
    tree = Tree("📄 文档结构", style="bold blue")
    for child in doc_structure.root.children:
        add_node_to_tree(tree, child)

    # 显示树
    console.print("\n")
    console.print(Panel(tree, border_style="blue"))
    console.print()


def display_comparison(new_models: Set[str], removed_models: Set[str]) -> None:
    """
    显示型号对比结果

    Args:
        new_models: 新增型号集合
        removed_models: 移除型号集合
    """
    # 创建对比表格
    compare_table = Table(
        title="🔄 型号对比",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    compare_table.add_column("变更类型", style="cyan")
    compare_table.add_column("数量", justify="right", style="green")
    compare_table.add_column("型号列表", style="yellow")

    # 添加新增型号
    if new_models:
        models_text = "\n".join(f"✨ {model}" for model in sorted(new_models))
        compare_table.add_row("➕ 新增", str(len(new_models)), models_text)

    # 添加移除型号
    if removed_models:
        models_text = "\n".join(f"❌ {model}" for model in sorted(removed_models))
        compare_table.add_row("➖ 移除", str(len(removed_models)), models_text)

    if new_models or removed_models:
        console.print()
        console.print(compare_table)
        console.print()
    else:
        console.print(Panel("[green]✅ 没有型号变更[/green]", border_style="green"))


def display_batch_verification(batch_results: Dict[str, Any]) -> None:
    """
    显示批次验证结果

    Args:
        batch_results: 批次验证结果字典
    """
    if not batch_results:
        console.print(
            Panel(
                "[yellow]⚠️ 没有批次数据可供验证[/yellow]",
                title="批次验证",
                border_style="yellow",
            )
        )
        return

    # 创建批次汇总表格
    summary_table = Table(
        title="🔍 批次数据汇总",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    summary_table.add_column("批次", style="cyan")
    summary_table.add_column("记录数", justify="right", style="green")
    summary_table.add_column("表格数", justify="right", style="yellow")

    # 计算批次总数，如果超过一定数量，只显示部分
    batch_count = len(batch_results)
    show_all = batch_count <= 50  # 只有50个批次以内才全部显示

    # 添加批次数据
    total_records = 0
    total_tables = 0

    sorted_batches = sorted(batch_results.items())

    # 如果批次太多，只显示前20个和后20个
    if not show_all:
        display_batches = sorted_batches[:20] + sorted_batches[-20:]
        console.print(
            f"[yellow]注意：只显示前20个和后20个批次（共{batch_count}个批次）[/yellow]"
        )
    else:
        display_batches = sorted_batches

    for batch, data in display_batches:
        total_records += data["total"]
        table_count = len(data["table_counts"])
        total_tables += table_count
        summary_table.add_row(f"第{batch}批", str(data["total"]), str(table_count))

    # 如果有省略的批次，添加省略提示行
    if not show_all and batch_count > 40:
        summary_table.add_row(f"... (省略 {batch_count - 40} 个批次) ...", "...", "...")

    # 添加合计行
    summary_table.add_row(
        "[bold]合计[/bold]",
        f"[bold]{total_records}[/bold]",
        f"[bold]{total_tables}[/bold]",
    )

    # 在表格前添加标题，表明这是关键信息
    console.print()
    console.print("[bold cyan]📊 关键信息：批次数据汇总[/bold cyan]")
    console.print(summary_table)
    console.print()


def display_consistency_result(result: Dict[str, Any]) -> None:
    """
    显示批次一致性验证结果

    Args:
        result: 一致性验证结果字典
    """
    # 在显示结果前添加标题，表明这是关键信息
    console.print()
    console.print("[bold cyan]📊 关键信息：数据一致性检查[/bold cyan]")

    if result["status"] == "no_batch":
        console.print(
            Panel(
                "[yellow]⚠️ 未找到批次号，无法验证数据一致性[/yellow]",
                title="数据一致性检查",
                border_style="yellow",
            )
        )
        return

    if result["status"] == "unknown":
        console.print(
            Panel(
                f"[yellow]⚠️ 第{result['batch']}批：未找到总记录数声明，实际记录数为 {result['actual_count']}[/yellow]",
                title="数据一致性检查",
                border_style="yellow",
            )
        )
    elif result["status"] == "match":
        console.print(
            Panel(
                f"[green]✅ 第{result['batch']}批：记录数匹配，共 {result['actual_count']} 条记录[/green]",
                title="数据一致性检查",
                border_style="green",
            )
        )
    elif result["status"] == "mismatch":
        diff_text = f"差异 {result['difference']} 条" if "difference" in result else ""
        console.print(
            Panel(
                f"[red]❌ 第{result['batch']}批：记录数不匹配！声明 {result['declared_count']}, 实际 {result['actual_count']}, {diff_text}[/red]",
                title="⚠️ 数据一致性检查",
                border_style="red",
            )
        )
    elif result["status"] == "internal_match":
        console.print(
            Panel(
                f"[green]✅ 第{result['batch']}批：内部一致性检查通过，表格记录总数 {result['actual_count']} 与处理结果数 {result['processed_count']} 一致[/green]",
                title="数据一致性检查",
                border_style="green",
            )
        )
    elif result["status"] == "internal_mismatch":
        diff_text = f"差异 {result['difference']} 条" if "difference" in result else ""
        console.print(
            Panel(
                f"[red]❌ 第{result['batch']}批：内部一致性检查失败！表格记录总数 {result['actual_count']} 与处理结果数 {result['processed_count']} 不一致，{diff_text}[/red]",
                title="⚠️ 数据一致性检查",
                border_style="red",
            )
        )

    # 显示表格记录分布
    table_counts = result.get("table_counts", {})
    if table_counts:
        count_table = Table(
            title="📊 表格记录分布",
            title_style="bold cyan",
            show_header=True,
            header_style="bold green",
            border_style="blue",
        )
        count_table.add_column("表格ID", style="cyan")
        count_table.add_column("记录数", justify="right", style="green")
        count_table.add_column("占比", justify="right", style="yellow")

        total = result.get("actual_count", sum(table_counts.values()))

        for table_id, count in sorted(table_counts.items()):
            percentage = (count / total * 100) if total > 0 else 0
            count_table.add_row(
                f"表格 {table_id}" if not isinstance(table_id, str) else table_id,
                str(count),
                f"{percentage:.1f}%",
            )

        console.print(count_table)


def create_progress_bar(total: int) -> Progress:
    """
    创建进度条并返回进度条对象

    Args:
        total: 任务总数

    Returns:
        Progress对象，已添加任务
    """
    progress = Progress(
        "[progress.description]{task.description}",
        SpinnerColumn(),
        BarColumn(bar_width=None),
        "[progress.percentage]{task.percentage:>3.1f}%",
        "•",
        "[bold blue]{task.completed}/{task.total}",
        "•",
        TimeRemainingColumn(),
        "•",
        TimeElapsedColumn(),
        console=console,
        transient=True,
    )

    progress.add_task("[bold cyan]🔄 处理文件", total=total)

    return progress


def print_docx_content(doc_path: str) -> None:
    """
    打印文档内容预览，显示所有元素的详细信息

    Args:
        doc_path: 文档路径
    """
    try:
        from docx import Document
        from docx.document import Document as DocxDocument

        doc: DocxDocument = Document(doc_path)
        console.print(
            Panel(
                f"[bold cyan]文件详细内容: {doc_path}[/bold cyan]", border_style="cyan"
            )
        )

        # 创建一个树形结构
        tree = Tree(f"📄 {doc_path}", style="bold blue")

        # 添加段落内容
        paragraphs_node = tree.add("[bold magenta]📝 段落内容[/bold magenta]")
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else "默认样式"
                para_node = paragraphs_node.add(
                    f"[blue]段落 {i}[/blue] ([yellow]{style_name}[/yellow])"
                )
                if "批" in text:
                    para_node.add(f"🔖 [bold red]{text}[/bold red]")
                elif "节能型汽车" in text or "新能源汽车" in text:
                    para_node.add(f"📌 [bold green]{text}[/bold green]")
                elif text.startswith("（") and not any(str.isdigit() for str in text):
                    para_node.add(f"📎 [bold yellow]{text}[/bold yellow]")
                elif any(
                    marker in text
                    for marker in ["勘误", "关于", "符合", "技术要求", "自动转入"]
                ):
                    para_node.add(f"ℹ️ [bold magenta]{text}[/bold magenta]")
                else:
                    para_node.add(Text(textwrap.shorten(text, width=100)))

        # 添加表格内容
        tables_node = tree.add("[bold cyan]📊 表格内容[/bold cyan]")
        for i, table in enumerate(doc.tables, 1):
            if table.rows:
                table_node = tables_node.add(
                    f"[blue]表格 {i}[/blue] ({len(table.rows)}行 x {len(table.rows[0].cells)}列)"
                )

                # 创建表格预览
                preview_table = Table(
                    show_header=True,
                    header_style="bold green",
                    border_style="blue",
                    title=f"表格 {i} 预览",
                    title_style="bold cyan",
                )

                # 添加表头
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                for header in headers:
                    preview_table.add_column(header, overflow="fold")

                # 添加数据行预览
                for row_idx, row in enumerate(table.rows[1:6], 1):  # 只显示前5行数据
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):  # 跳过空行
                        preview_table.add_row(*cells)

                table_node.add(preview_table)

        console.print(tree)

    except Exception as e:
        console.print(
            Panel(
                f"[bold red]预览文件 {doc_path} 时出错: {e}[/bold red]",
                border_style="red",
            )
        )


def generate_ascii_bar_chart(
    data: Dict[str, int], title: str, width: int = 40
) -> Panel:
    """
    生成ASCII文本形式的柱状图

    Args:
        data: 数据字典，键为标签，值为数值
        title: 图表标题
        width: 图表最大宽度

    Returns:
        包装在Panel中的图表
    """
    # 确定最大值和标签长度
    max_value = max(data.values()) if data else 0
    max_label_length = max(len(label) for label in data.keys()) if data else 0

    if max_value == 0:
        return Panel(
            f"[yellow]没有数据可显示[/yellow]", title=title, border_style="blue"
        )

    # 生成柱状图
    chart_lines = []
    chart_lines.append(f"[bold cyan]{title}[/bold cyan]")
    chart_lines.append("")

    for label, value in sorted(data.items(), key=lambda x: x[1], reverse=True):
        # 计算柱长度
        bar_length = int((value / max_value) * width)
        bar = "█" * bar_length

        # 格式化输出
        percentage = (value / sum(data.values())) * 100
        chart_lines.append(
            f"{label.ljust(max_label_length)} │ {bar} {value} ({percentage:.1f}%)"
        )

    chart_text = "\n".join(chart_lines)
    return Panel(chart_text, border_style="blue", box=box.ROUNDED)


def generate_spark_line(data: List[int], title: str, width: int = 40) -> Panel:
    """
    生成简单的spark line图表

    Args:
        data: 数据点列表
        title: 图表标题
        width: 图表宽度

    Returns:
        包装在Panel中的图表
    """
    if not data or all(x == 0 for x in data):
        return Panel(
            f"[yellow]没有数据可显示[/yellow]", title=title, border_style="blue"
        )

    # 缩放数据到0-7的范围（使用Unicode方块字符的8个高度级别）
    min_val = min(data)
    max_val = max(data)
    range_val = max_val - min_val if max_val > min_val else 1

    # 使用Unicode方块字符表示不同高度
    spark_chars = "▁▂▃▄▅▆▇█"

    # 生成spark line
    if len(data) > width:
        # 如果数据点太多，需要采样
        step = len(data) / width
        sampled_data = [data[int(i * step)] for i in range(width)]
    else:
        # 如果数据点不够，进行填充
        sampled_data = data + [data[-1]] * (width - len(data)) if data else []

    # 生成spark line字符
    spark_line = ""
    for val in sampled_data:
        if val == min_val:
            spark_line += spark_chars[0]
        elif val == max_val:
            spark_line += spark_chars[-1]
        else:
            index = int(((val - min_val) / range_val) * (len(spark_chars) - 1))
            spark_line += spark_chars[index]

    # 构建图表文本
    chart_text = f"[bold cyan]{title}[/bold cyan]\n\n"
    chart_text += f"{spark_line}\n"
    chart_text += (
        f"最小值: {min_val}  最大值: {max_val}  平均值: {sum(data)/len(data):.1f}"
    )

    return Panel(chart_text, border_style="blue", box=box.ROUNDED)


def display_summary_dashboard(
    cars_data: List[Dict[str, Any]],
    batch_results: Dict[str, Any],
    consistency_result: Dict[str, Any],
    output_file: str,
) -> None:
    """
    显示处理结果汇总面板，将各种统计和验证结果整合到一个统一的仪表盘

    Args:
        cars_data: 车辆数据列表
        batch_results: 批次验证结果
        consistency_result: 一致性检查结果
        output_file: 输出文件路径
    """
    from ..batch.validator import calculate_statistics

    # 计算统计信息
    stats = calculate_statistics(cars_data)
    total_count = stats["total_count"]
    energy_saving_count = stats["energy_saving_count"]
    new_energy_count = stats["new_energy_count"]

    # 创建车辆类型分布图
    vehicle_type_data = {
        "节能型汽车": energy_saving_count,
        "新能源汽车": new_energy_count,
    }
    type_chart = generate_ascii_bar_chart(vehicle_type_data, "车辆类型分布")

    # 创建布局
    layout = Layout(name="dashboard")
    layout.split(
        Layout(name="header", size=3),
        Layout(name="main", ratio=1),
        Layout(name="footer", size=3),
    )

    layout["main"].split_row(
        Layout(name="left", ratio=1),
        Layout(name="right", ratio=2),  # 给右侧更多空间
    )

    # 创建标题
    title_text = Text("📊 车辆数据处理结果汇总", style="bold white on blue")
    title_text = Text.assemble(
        title_text, Text(f" | 共处理 {total_count} 条记录", style="bold white")
    )

    # 创建统计表格
    stats_table = Table(
        title="数据统计",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
        box=box.ROUNDED,
    )

    # 添加列
    stats_table.add_column("类型", style="cyan")
    stats_table.add_column("数量", justify="right", style="green")
    stats_table.add_column("占比", justify="right", style="yellow")

    # 计算百分比
    energy_saving_percent = (
        energy_saving_count / total_count * 100 if total_count > 0 else 0
    )
    new_energy_percent = new_energy_count / total_count * 100 if total_count > 0 else 0

    # 添加行
    stats_table.add_row(
        "🚗 节能型汽车", f"{energy_saving_count:,}", f"{energy_saving_percent:.1f}%"
    )
    stats_table.add_row(
        "⚡ 新能源汽车", f"{new_energy_count:,}", f"{new_energy_percent:.1f}%"
    )
    stats_table.add_row("📝 总记录数", f"{total_count:,}", "100%")

    # 创建批次分布表格
    batch_count_table = Table(
        title="批次分布",
        show_header=True,
        header_style="bold green",
        title_style="bold cyan",
        border_style="blue",
        box=box.ROUNDED,
    )

    batch_count_table.add_column("批次", style="cyan")
    batch_count_table.add_column("数量", justify="right", style="green")
    batch_count_table.add_column("占比", justify="right", style="yellow")

    # 添加批次数据
    batch_counts = stats.get("batch_counts", {})
    sorted_batches = sorted(batch_counts.items())

    # 决定显示多少批次（基于可用空间）
    display_count = min(20, len(sorted_batches))  # 默认最多显示20个批次

    # 为批次分布图准备数据
    batch_chart_data = {}

    for batch, count in sorted_batches[:display_count]:
        percentage = (count / total_count) * 100
        batch_count_table.add_row(f"第{batch}批", f"{count:,}", f"{percentage:.1f}%")

        # 只取前10个批次用于图表显示
        if len(batch_chart_data) < 10:
            batch_chart_data[f"第{batch}批"] = count

    if len(batch_counts) > display_count:
        remaining_count = sum(count for _, count in sorted_batches[display_count:])
        remaining_percentage = (remaining_count / total_count) * 100
        batch_count_table.add_row(
            f"其他批次 (共{len(batch_counts) - display_count}个)",
            f"{remaining_count:,}",
            f"{remaining_percentage:.1f}%",
        )

        # 如果批次太多，添加"其他"类别到图表
        if len(sorted_batches) > 10:
            other_count = sum(count for _, count in sorted_batches[10:])
            batch_chart_data["其他批次"] = other_count

    # 添加合计行
    batch_count_table.add_row(
        "[bold]合计[/bold]", f"[bold]{total_count}[/bold]", f"[bold]100%[/bold]"
    )

    # 创建批次分布图
    batch_chart = generate_ascii_bar_chart(batch_chart_data, "批次分布图表")

    # 创建一致性状态面板，同时包含输出信息
    if consistency_result["status"] in ["match", "internal_match"]:
        status_style = "green"
        status_icon = "✅"
        status_text = "数据一致"
    elif consistency_result["status"] in ["mismatch", "internal_mismatch"]:
        status_style = "red"
        status_icon = "❌"
        status_text = "数据不一致"
    else:
        status_style = "yellow"
        status_icon = "⚠️"
        status_text = "未知状态"

    # 合并一致性检查和输出信息到一个面板
    info_panel = Panel(
        f"[{status_style}]{status_icon} 一致性检查: {status_text}[/{status_style}]\n"
        f"批次: 第{consistency_result.get('batch', '未知')}批\n"
        f"实际记录: {consistency_result.get('actual_count', '未知')}\n"
        f"期望记录: {consistency_result.get('declared_count', consistency_result.get('processed_count', '未知'))}\n\n"
        f"[blue]📂 输出文件:[/blue] {output_file}\n"
        f"[blue]🕒 处理完成时间:[/blue] {time.strftime('%Y-%m-%d %H:%M:%S')}",
        title="处理信息",
        border_style="blue",
        box=box.ROUNDED,
    )

    # 组装左侧布局
    left_content = Layout()
    left_content.split(
        Layout(stats_table, name="stats", ratio=1),
        Layout(type_chart, name="chart", ratio=1),
        Layout(info_panel, name="info", ratio=1),
    )

    # 组装右侧布局 - 根据批次数量决定布局
    right_content = Layout()
    if len(batch_counts) > 5:  # 如果批次数量较多，添加图表
        right_content.split(
            Layout(batch_count_table, name="batch_table", ratio=2),
            Layout(batch_chart, name="batch_chart", ratio=1),
        )
        layout["right"].update(right_content)
    else:
        # 批次少时直接显示表格
        layout["right"].update(batch_count_table)

    # 组装布局
    layout["header"].update(Panel(title_text, border_style="blue", box=box.ROUNDED))
    layout["left"].update(left_content)

    footer_text = Text(
        "💡 使用 -v 参数查看更详细的信息 | 🔍 对比过往批次 | 📥 查看更多统计数据",
        style="bold white on blue",
    )
    layout["footer"].update(Panel(footer_text, border_style="blue", box=box.ROUNDED))

    # 显示布局
    console.print()
    console.print(layout)
    console.print()
