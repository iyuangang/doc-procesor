"""
文档处理器模块 - 提供文档处理的核心功能
"""

import gc
import logging
import os
import re
import time
import tempfile
import shutil
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple, Set

import psutil
from docx import Document
from docx.document import Document as DocxDocument
import pandas as pd

from ..batch.validator import verify_batch_consistency, verify_all_batches
from ..config.settings import settings
from ..document.parser import (
    extract_doc_content,
    extract_declared_count,
    get_table_type,
)
from ..models.document_node import DocumentNode, DocumentStructure
from ..table.extractor import TableExtractor
from ..utils.chinese_numbers import extract_batch_number
from ..utils.validation import process_car_info
from ..ui.console import (
    display_consistency_result,
    display_doc_content,
    display_batch_verification,
    display_statistics,
)


class ProcessingError(Exception):
    """处理错误异常"""

    pass


class DocumentError(Exception):
    """文档错误异常"""

    pass


class DocProcessor:
    """文档处理器类，用于处理Word文档中的车辆信息"""

    def __init__(
        self,
        doc_path: str,
        verbose: bool = True,
        config: Optional[Dict[str, Any]] = None,
    ):
        """
        初始化文档处理器

        Args:
            doc_path: 文档路径
            verbose: 是否显示详细信息
            config: 配置信息
        """
        self.doc_path = doc_path
        self.start_time = time.time()
        self.config = config or {}
        self.logger = logging.getLogger(__name__)
        self.doc_structure = DocumentStructure()

        try:
            self._load_document()
        except Exception as e:
            self.logger.error(f"初始化文档处理器失败: {str(e)}")
            raise DocumentError(f"无法加载文档 {doc_path}: {str(e)}")

        self.current_category: Optional[str] = None
        self.current_type: Optional[str] = None
        self.batch_number: Optional[str] = None
        self.cars: List[Dict[str, Any]] = []
        self._processing_times: Dict[str, float] = {}
        self.declared_count: Optional[int] = None  # 声明的总记录数

        # 从配置加载设置
        self._chunk_size = self._get_config("performance.chunk_size", 1000)
        self.verbose = verbose
        self._cache_size_limit = self._get_config(
            "performance.cache_size_limit", 50 * 1024 * 1024
        )
        self._cleanup_interval = self._get_config("performance.cleanup_interval", 300)
        # 添加跳过总记录数检查的配置选项
        self._skip_count_check = self._get_config("document.skip_count_check", False)
        # 设置搜索限制
        self._max_paragraphs_to_search = self._get_config(
            "document.max_paragraphs_to_search", 30
        )
        self._max_tables_to_search = self._get_config(
            "document.max_tables_to_search", 5
        )

        # 预编译正则表达式
        self._batch_pattern = re.compile(r"第([一二三四五六七八九十百零\d]+)批")
        self._whitespace_pattern = re.compile(r"\s+")
        self._chinese_number_pattern = re.compile(r"([一二三四五六七八九十百零]+)")
        self._count_pattern = re.compile(
            r"(共计|总计|合计).*?(\d+).*?(款|个|种|辆|台|项)"
        )  # 总记录数模式

        self._last_cache_cleanup = time.time()
        self.logger.info(f"初始化文档处理器: {doc_path}")

        self.current_section: Optional[DocumentNode] = None
        self.current_subsection: Optional[DocumentNode] = None
        self.current_numbered_section: Optional[DocumentNode] = (
            None  # 用于跟踪带数字编号的节点
        )

        # 初始化表格提取器
        self.table_extractor = TableExtractor(chunk_size=self._chunk_size)

    def _get_config(self, key: str, default: Any) -> Any:
        """
        获取配置值

        Args:
            key: 配置键，使用点号分隔，例如 'performance.chunk_size'
            default: 默认值

        Returns:
            配置值
        """
        if self.config:
            # 处理嵌套键
            if "." in key:
                parts = key.split(".")
                value = self.config
                for part in parts:
                    if isinstance(value, dict) and part in value:
                        value = value[part]
                    else:
                        return default
                return value
            return self.config.get(key, default)

        # 如果没有提供配置，尝试从全局设置获取
        try:
            return settings.get(key, default)
        except:
            return default

    def _load_document(self) -> None:
        """
        安全加载文档，处理大文件

        Raises:
            DocumentError: 无法加载文档
        """
        try:
            file_size = os.path.getsize(self.doc_path)
            self.logger.info(
                f"加载文档 {self.doc_path}, 大小: {file_size / 1024 / 1024:.2f}MB"
            )

            large_file_threshold = (
                self._get_config("document.large_file_threshold", 100) * 1024 * 1024
            )
            if file_size > large_file_threshold:  # 配置的阈值，默认100MB
                self.logger.warning(
                    f"文档大小超过{large_file_threshold / 1024 / 1024}MB, 使用临时文件处理"
                )
                with tempfile.NamedTemporaryFile(delete=False) as tmp:
                    shutil.copy2(self.doc_path, tmp.name)
                    self.doc = Document(tmp.name)
                    os.unlink(tmp.name)
            else:
                self.doc = Document(self.doc_path)
        except Exception as e:
            self.logger.error(f"加载文档失败: {str(e)}")
            raise DocumentError(f"无法加载文档 {self.doc_path}: {str(e)}")

    def _log_time(self, operation: str) -> None:
        """
        记录操作耗时

        Args:
            operation: 操作名称
        """
        current_time = time.time()
        elapsed = current_time - self.start_time
        self._processing_times[operation] = elapsed
        if operation != "init" and self.verbose:
            self.logger.debug(f"{operation} 耗时: {elapsed:.2f}秒")
        self.start_time = current_time

    def _extract_declared_count(self) -> Optional[int]:
        """
        从文档中提取批次声明的总记录数

        Returns:
            声明的总记录数，如果未找到则返回None
        """
        # 如果配置了跳过总记录数检查，直接返回None
        if self._skip_count_check:
            self.logger.info("根据配置跳过总记录数检查")
            return None

        return extract_declared_count(
            self.doc_path,
            max_paragraphs=self._max_paragraphs_to_search,
            max_tables=self._max_tables_to_search,
        )

    def process(self) -> List[Dict[str, Any]]:
        """
        处理文档并返回所有车辆信息

        Returns:
            车辆信息字典列表

        Raises:
            ProcessingError: 处理文档失败
        """
        try:
            self.logger.info(f"开始处理文档: {self.doc_path}")
            self._log_time("init")

            table_count = 0
            row_count = 0
            error_count = 0

            # 遍历文档中的所有元素
            for element in self.doc.element.body:
                try:
                    # 处理段落
                    if element.tag.endswith("p"):
                        text = element.text.strip()
                        if not text:
                            continue

                        # 提取批次号
                        if not self.batch_number:
                            self.batch_number = extract_batch_number(text)
                            if self.batch_number:
                                self.doc_structure.set_batch_number(self.batch_number)
                                self.logger.info(f"提取到批次号: {self.batch_number}")
                                self.doc_structure.add_node(
                                    f"第{self.batch_number}批", "batch", level=0
                                )

                        # 更新分类信息
                        if "节能型汽车" in text:
                            self.current_category = "节能型"
                            self.current_section = self.doc_structure.add_node(
                                "节能型汽车", "section", content=text
                            )
                            self.current_subsection = None
                            self.current_numbered_section = None
                            self.logger.debug(f"更新分类: {self.current_category}")
                        elif "新能源汽车" in text:
                            self.current_category = "新能源"
                            self.current_section = self.doc_structure.add_node(
                                "新能源汽车", "section", content=text
                            )
                            self.current_subsection = None
                            self.current_numbered_section = None
                            self.logger.debug(f"更新分类: {self.current_category}")
                        elif text.startswith("（") and not any(
                            str.isdigit() for str in text
                        ):
                            self.current_subsection = self.doc_structure.add_node(
                                text.strip(),
                                "subsection",
                                content=text,
                                parent_node=self.current_section,
                            )
                            self.current_numbered_section = None
                            self.logger.debug(f"更新类型: {text}")
                        # 处理带数字编号的节点
                        elif text.startswith(("1.", "2.", "3.", "4.", "5.")):
                            self.current_numbered_section = self.doc_structure.add_node(
                                text.strip(),
                                "numbered_section",
                                content=text,
                                parent_node=self.current_subsection
                                or self.current_section,
                            )
                            self.logger.debug(f"更新编号节点: {text}")
                        # 处理带括号数字编号的子节点
                        elif text.startswith("（") and any(
                            num in text for num in "123456789"
                        ):
                            if self.current_numbered_section:
                                self.doc_structure.add_node(
                                    text.strip(),
                                    "numbered_subsection",
                                    content=text,
                                    parent_node=self.current_numbered_section,
                                )
                            else:
                                self.doc_structure.add_node(
                                    text.strip(),
                                    "numbered_subsection",
                                    content=text,
                                    parent_node=self.current_subsection
                                    or self.current_section,
                                )
                            self.logger.debug(f"更新编号子节点: {text}")
                        elif "勘误" in text or "说明" in text:
                            self.doc_structure.add_node(
                                text[:40] + "...",
                                "note",
                                content=text,
                                parent_node=self.current_section,
                            )
                        elif "更正" in text or "修改" in text:
                            self.doc_structure.add_node(
                                text[:40] + "...",
                                "correction",
                                content=text,
                                parent_node=self.current_section,
                            )
                        else:
                            self.doc_structure.add_node(
                                text[:40] + "...",
                                "text",
                                content=text,
                                parent_node=self.current_section,
                            )

                    # 处理表格
                    elif element.tag.endswith("tbl"):
                        table_count += 1
                        for i, table in enumerate(self.doc.tables):
                            if table._element is element:
                                if table.rows:
                                    row_count += len(table.rows)
                                try:
                                    # 确定表格所属的类别和子类型
                                    current_sub_type = (
                                        self.current_subsection.title
                                        if self.current_subsection
                                        else None
                                    )

                                    # 提取表格中的车辆信息
                                    table_cars = self.table_extractor.extract_car_info(
                                        table,
                                        i,
                                        self.current_category,
                                        current_sub_type,
                                        self.batch_number,
                                    )
                                    self.cars.extend(table_cars)

                                    # 添加表格节点到正确的父节点
                                    parent_node = (
                                        self.current_numbered_section
                                        or self.current_subsection
                                        or self.current_section
                                    )
                                    self.doc_structure.add_node(
                                        f"表格 {i + 1}",
                                        "table",
                                        metadata={
                                            "rows": len(table.rows),
                                            "columns": len(table.rows[0].cells)
                                            if table.rows
                                            else 0,
                                            "records": len(table_cars),
                                            "category": self.current_category,
                                            "sub_type": current_sub_type,
                                        },
                                        parent_node=parent_node,
                                    )

                                    if self.verbose:
                                        self.logger.info(
                                            f"处理表格 {i + 1}, 提取到 {len(table_cars)} 条记录"
                                        )
                                except Exception as e:
                                    error_count += 1
                                    self.logger.error(
                                        f"处理表格 {i + 1} 出错: {str(e)}"
                                    )
                                break
                except Exception as e:
                    error_count += 1
                    self.logger.error(f"处理元素出错: {str(e)}")
                    continue

            self._log_time("process")
            self.logger.info(
                f"文档处理完成: {table_count} 个表格, {row_count} 行, "
                f"{len(self.cars)} 条记录, {error_count} 个错误"
            )

            # 执行批次数据一致性验证 - 只在处理后执行一次
            verification_start = time.time()

            # 获取声明的总记录数（如果尚未获取）
            if self.declared_count is None:
                self.declared_count = self._extract_declared_count()

            consistency_result = verify_batch_consistency(
                self.cars, self.batch_number, self.declared_count
            )
            verification_time = time.time() - verification_start
            self.logger.info(
                f"批次一致性验证结果: {consistency_result['status']} (耗时: {verification_time:.2f}秒)"
            )

            # 计算文件大小和记录数
            file_size = os.path.getsize(self.doc_path) / (1024 * 1024)  # MB
            record_count = len(self.cars)

            # 对于大文件或大量记录，禁用详细显示以提高性能
            is_large_file = file_size > 50 or record_count > 100000  # 50MB或10万条记录

            # 显示文档结构（仅在详细模式下）
            if self.verbose and not is_large_file:
                display_doc_content(self.doc_structure)
            elif self.verbose and is_large_file:
                self.logger.info("文件较大，跳过显示详细文档结构以提高性能")

            # 显示批次验证结果
            batch_results = verify_all_batches(self.cars)
            if (
                self._get_config("output.show_key_info_in_compact_mode", True)
                or self.verbose
            ) and batch_results:
                display_batch_verification(batch_results)

            # 显示批次一致性验证结果（始终显示，即使在简洁模式下）
            if (
                self._get_config("output.show_key_info_in_compact_mode", True)
                or self.verbose
            ):
                display_consistency_result(consistency_result)

            # 处理完成后主动释放资源
            self.table_extractor.clear_cache()
            gc.collect()

            return self.cars

        except Exception as e:
            self.logger.error(f"处理文档失败: {str(e)}")
            raise ProcessingError(f"处理文档 {self.doc_path} 失败: {str(e)}")

    def get_memory_usage(self) -> str:
        """
        获取当前进程的内存使用情况

        Returns:
            内存使用情况字符串
        """
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        return f"{memory_info.rss / 1024 / 1024:.1f}MB"

    def save_to_csv(self, output_file: str) -> None:
        """
        将处理结果保存为CSV文件

        Args:
            output_file: 输出文件路径
        """
        if not self.cars:
            self.logger.warning("没有数据可保存")
            return

        # 估计数据大小
        estimated_size = len(self.cars) * 500  # 假设每条记录约500字节
        is_large_dataset = estimated_size > 100 * 1024 * 1024  # 100MB

        if is_large_dataset:
            self.logger.info(f"大数据集 ({len(self.cars)} 条记录), 使用优化处理...")

            # 使用分块处理
            chunk_size = self._chunk_size
            with open(output_file, "w", encoding="utf-8-sig") as f:
                # 写入表头
                first_batch = self.cars[:100]  # 取前100条确定字段
                all_fields: Set[str] = set()
                for car in first_batch:
                    all_fields.update(car.keys())

                base_columns = [
                    "batch",
                    "energytype",
                    "vmodel",
                    "category",
                    "sub_type",
                    "序号",
                    "企业名称",
                    "品牌",
                    "table_id",
                    "raw_text",
                ]

                header_fields = [col for col in base_columns if col in all_fields] + [
                    col for col in sorted(all_fields) if col not in base_columns
                ]

                f.write(",".join(header_fields) + "\n")

                # 分块写入数据
                for i in range(0, len(self.cars), chunk_size):
                    chunk = self.cars[i : i + chunk_size]
                    chunk_df = pd.DataFrame(chunk)
                    chunk_df = chunk_df.reindex(columns=header_fields)

                    if i == 0:
                        chunk_df.to_csv(
                            f, index=False, header=False, encoding="utf-8-sig"
                        )
                    else:
                        chunk_df.to_csv(
                            f, index=False, header=False, encoding="utf-8-sig", mode="a"
                        )

                    # 释放内存
                    del chunk_df
                    gc.collect()

            self.logger.info(f"保存完成, 文件: {output_file}, 记录数: {len(self.cars)}")
        else:
            # 原有处理逻辑
            all_cars_df = pd.DataFrame(self.cars)

            # 优化列顺序设置
            base_columns = [
                "batch",
                "energytype",
                "vmodel",
                "category",
                "sub_type",
                "序号",
                "企业名称",
                "品牌",
                "table_id",
                "raw_text",
            ]
            all_columns = all_cars_df.columns.tolist()
            final_columns = [col for col in base_columns if col in all_columns] + [
                col for col in all_columns if col not in base_columns
            ]

            # 重新排列列并保存
            all_cars_df = all_cars_df[final_columns]
            all_cars_df.to_csv(output_file, index=False, encoding="utf-8-sig")

            self.logger.info(
                f"保存完成, 文件: {output_file}, 记录数: {len(all_cars_df)}"
            )


def process_doc(
    doc_path: str,
    output_file: Optional[str] = None,
    verbose: bool = False,
    config: Optional[Dict[str, Any]] = None,
) -> List[Dict[str, Any]]:
    """
    处理单个文档的函数

    Args:
        doc_path: 文档路径
        output_file: 输出文件路径，如果提供则保存CSV文件
        verbose: 是否显示详细信息
        config: 配置信息

    Returns:
        车辆信息字典列表
    """
    try:
        processor = DocProcessor(doc_path, verbose, config)
        result = processor.process()

        # 如果提供了输出文件路径，则保存CSV文件
        if output_file and result:
            processor.save_to_csv(output_file)

            # 显示批次验证结果
            batch_results = verify_all_batches(result)
            if batch_results:
                display_batch_verification(batch_results)

            # 显示统计信息
            if verbose and result:
                from ..batch.validator import calculate_statistics

                stats = calculate_statistics(result)
                display_statistics(
                    stats["total_count"],
                    stats["energy_saving_count"],
                    stats["new_energy_count"],
                    output_file,
                )

        return result
    except Exception as e:
        logging.error(f"处理文档 {doc_path} 失败: {str(e)}")
        return []
