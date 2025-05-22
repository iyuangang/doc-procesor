"""
文档解析模块 - 提供从Word文档中提取内容的功能
"""

import re
from typing import List, Dict, Tuple, Any, Optional

from docx import Document
from docx.document import Document as DocxDocument

from ..utils.chinese_numbers import extract_batch_number
from ..utils.text_processing import clean_text


def extract_doc_content(doc_path: str) -> Tuple[List[str], List[Dict[str, str]]]:
    """
    提取文档中除表格外的内容，并分离额外信息

    Args:
        doc_path: 文档路径

    Returns:
        元组，包含段落列表和额外信息列表
    """
    doc: DocxDocument = Document(doc_path)
    paragraphs: List[str] = []
    extra_info: List[Dict[str, str]] = []
    current_section: Optional[str] = None
    batch_found = False
    batch_number = None

    # 额外信息的标识词和对应类型
    info_types: Dict[str, str] = {
        "勘误": "勘误",
        "关于": "政策",
        "符合": "说明",
        "技术要求": "说明",
        "自动转入": "说明",
        "第二部分": "说明",
    }

    # 用于收集连续的额外信息文本
    current_extra_info: Optional[Dict[str, str]] = None

    def save_current_extra_info() -> None:
        """保存当前的额外信息"""
        nonlocal current_extra_info
        if current_extra_info:
            # 清理和规范化内容
            content = current_extra_info["content"]
            # 移除多余的空白字符
            content = re.sub(r"\s+", " ", content)
            # 移除换行符
            content = content.replace("\n", " ")
            current_extra_info["content"] = content.strip()

            # 添加批次号
            if batch_number:
                current_extra_info["batch"] = batch_number

            # 检查是否需要合并相同类型和章节的信息
            for info in extra_info:
                if (
                    info["type"] == current_extra_info["type"]
                    and info["section"] == current_extra_info["section"]
                ):
                    info["content"] = (
                        info["content"] + " " + current_extra_info["content"]
                    )
                    current_extra_info = None
                    return

            extra_info.append(current_extra_info)
            current_extra_info = None

    # 遍历文档段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # 如果遇到空行，保存当前的额外信息
            if current_extra_info:
                save_current_extra_info()
            continue

        # 检查批次号
        if not batch_found and "批" in text:
            extracted_batch = extract_batch_number(text)
            if extracted_batch:
                batch_number = extracted_batch
                paragraphs.append(text)  # 将批次号信息放在最前面
                batch_found = True
                continue

        # 识别主要分类
        if "节能型汽车" in text or "新能源汽车" in text:
            save_current_extra_info()
            current_section = text
            paragraphs.append(text)
        # 识别子分类，排除括号中有数字的
        elif text.startswith("（") and not any(str.isdigit() for str in text):
            save_current_extra_info()
            current_section = text
            paragraphs.append(text)
        # 识别额外信息
        elif any(marker in text for marker in info_types.keys()):
            # 如果当前文本包含新的标识词，保存之前的信息并创建新的
            if current_extra_info:
                save_current_extra_info()

            # 创建新的额外信息
            info_type = next((t for m, t in info_types.items() if m in text), "其他")
            current_extra_info = {
                "section": current_section or "文档说明",
                "type": info_type,
                "content": text,
            }
        # 如果当前有未处理的额外信息，将文本追加到内容中
        elif current_extra_info is not None:
            current_extra_info["content"] = current_extra_info["content"] + " " + text
        else:
            paragraphs.append(text)

    # 保存最后一条未处理的额外信息
    save_current_extra_info()

    return paragraphs, extra_info


def get_table_type(
    headers: List[str], current_category: Optional[str], current_type: Optional[str]
) -> Tuple[str, str]:
    """
    根据表头判断表格类型，使用当前上下文确定子类型

    Args:
        headers: 表头列表
        current_category: 当前分类（节能型或新能源）
        current_type: 当前子类型（表格所在的子分类）

    Returns:
        (category, sub_type)元组
    """
    # 标准化表头
    normalized_headers = [h.strip().lower() for h in headers]

    # 验证必要的列是否存在
    required_columns = {"序号", "企业名称"}
    missing_columns = required_columns - set(normalized_headers)
    if missing_columns:
        raise ValueError(f"表格缺少必要的列: {missing_columns}")

    # 处理特殊的表头组合
    if "型式" in normalized_headers and "档位数" in normalized_headers:
        # 合并为变速器列
        idx = normalized_headers.index("型式")
        normalized_headers[idx] = "变速器"
        normalized_headers.pop(idx + 1)

    # 只从表头判断category（节能型或新能源）
    category = current_category or "未知"

    # 如果在明确的节能型部分中，优先使用节能型分类
    if "节能型" in str(current_category).lower():
        category = "节能型"

    # 如果在明确的新能源部分中，优先使用新能源分类
    if "新能源" in str(current_category).lower():
        category = "新能源"

    # 始终使用当前上下文的子类型，不从表头判断
    sub_type = current_type or "未知"

    # 确保返回的是字符串类型
    return str(category), str(sub_type)


def extract_declared_count(
    doc_path: str, max_paragraphs: int = 30, max_tables: int = 5
) -> Optional[int]:
    """
    从文档中提取批次声明的总记录数

    Args:
        doc_path: 文档路径
        max_paragraphs: 最多搜索的段落数
        max_tables: 最多搜索的表格数

    Returns:
        声明的总记录数，如果未找到则返回None
    """
    # 预编译正则表达式
    count_pattern = re.compile(r"(共计|总计|合计).*?(\d+).*?(款|个|种|辆|台|项)")

    try:
        doc: DocxDocument = Document(doc_path)

        # 1. 只搜索前N个段落
        paragraphs_to_search = min(max_paragraphs, len(doc.paragraphs))

        for para in doc.paragraphs[:paragraphs_to_search]:
            text = para.text.strip()
            if not text:
                continue

            if "总" in text or "共" in text or "合计" in text:
                match = count_pattern.search(text)
                if match:
                    try:
                        count = int(match.group(2))
                        return count
                    except (ValueError, IndexError):
                        continue

        # 2. 只搜索前M个表格
        tables_to_search = min(max_tables, len(doc.tables))

        for table in doc.tables[:tables_to_search]:
            if not table.rows:
                continue

            # 只检查表格的前3行和后3行，这些位置最可能出现合计信息
            rows_to_check = []
            if len(table.rows) > 6:
                rows_to_check = list(table.rows[:3]) + list(table.rows[-3:])
            else:
                rows_to_check = list(table.rows)

            for row in rows_to_check:
                cells = [cell.text.strip() for cell in row.cells]
                # 检查是否包含合计相关的内容
                if any(cell.startswith(("合计", "总计")) for cell in cells):
                    # 尝试从合计行中获取数值
                    for cell in cells:
                        if cell.isdigit():
                            return int(cell)

        return None

    except Exception as e:
        print(f"提取总记录数时出错: {e}")
        return None
