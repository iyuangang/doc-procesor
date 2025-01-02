import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx.document import Document
from docx.table import Table
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    import pandas as pd
else:
    pd = Mock()  # type: ignore

from main import (
    cn_to_arabic,
    extract_batch_number,
    clean_text,
    validate_car_info,
    get_table_type,
    process_car_info,
    extract_doc_content,
    DocProcessor,
)


# 测试中文数字转阿拉伯数字
@pytest.mark.parametrize(
    "input_num,expected",
    [
        ("一", "1"),
        ("二", "2"),
        ("十", "10"),
        ("十一", "11"),
        ("二十", "20"),
        ("二十一", "21"),
        ("一百", "100"),
        ("一百零一", "101"),
        ("123", "123"),  # 已经是阿拉伯数字
        ("abc", "abc"),  # 非数字
    ],
)
def test_cn_to_arabic(input_num: str, expected: str):
    assert cn_to_arabic(input_num) == expected


# 测试批次号提取
@pytest.mark.parametrize(
    "input_text,expected",
    [
        ("第六十五批", "65"),
        ("第一百零一批", "101"),
        ("第123批", "123"),
        ("第一批", "1"),
        ("无批次信息", None),
        ("", None),
    ],
)
def test_extract_batch_number(input_text: str, expected: str | None):
    assert extract_batch_number(input_text) == expected


# 测试文本清理
@pytest.mark.parametrize(
    "input_text,expected",
    [
        ("  测试文本  ", "测试文本"),
        ("测试，文本", "测试,文本"),
        ("测试；文本", "测试;文本"),
        ("测试\n文本", "测试 文本"),
        ("测试    文本", "测试 文本"),
    ],
)
def test_clean_text(input_text: str, expected: str):
    assert clean_text(input_text) == expected


# 测试车辆信息验证
@pytest.mark.parametrize(
    "car_info,expected_valid,expected_message",
    [
        (
            {"企业名称": "测试企业", "型号": "TEST001", "car_type": 1},
            True,
            "",
        ),
        (
            {"企业名称": "", "型号": "TEST001", "car_type": 1},
            False,
            "缺少必要字段: 企业名称",
        ),
        (
            {"企业名称": "测试企业", "型号": "", "car_type": 1},
            False,
            "缺少必要字段: 型号",
        ),
        (
            {"企业名称": "测试企业", "型号": "TEST001"},
            False,
            "缺少车型标识",
        ),
        (
            {"企业名称": "测试企业", "型号": "TEST001", "car_type": 3},
            False,
            "无效的车型标识: 3",
        ),
    ],
)
def test_validate_car_info(car_info: dict, expected_valid: bool, expected_message: str):
    valid, message = validate_car_info(car_info)
    assert valid == expected_valid
    assert message == expected_message


# 测试表格类型判断
@pytest.mark.parametrize(
    "headers,current_category,current_type,expected",
    [
        (
            ["排量(ml)", "综合燃料消耗量"],
            "节能型",
            "（一）乘用车",
            ("节能型", "（一）乘用车"),
        ),
        (
            ["燃料种类", "CNG"],
            "节能型",
            None,
            ("节能型", "（二）轻型商用车"),
        ),
        (
            ["燃料种类", "LNG"],
            "节能型",
            None,
            ("节能型", "（三）重型商用车"),
        ),
        (
            ["纯电动续驶里程", "燃料消耗量", "通用名称"],
            "新能源",
            None,
            ("新能源", "（一）插电式混合动力乘用车"),
        ),
        (
            ["纯电动续驶里程", "动力蓄电池总能量"],
            "新能源",
            None,
            ("新能源", "（二）纯电动商用车"),
        ),
        (
            ["纯电动续驶里程", "燃料消耗量"],
            "新能源",
            None,
            ("新能源", "（三）插电式混合动力商用车"),
        ),
        (
            ["燃料电池系统额定功率"],
            "新能源",
            None,
            ("新能源", "（四）燃料电池商用车"),
        ),
        (
            ["未知字段"],
            "未知",
            "未知",
            ("未知", "未知"),
        ),
    ],
)
def test_get_table_type(
    headers: list[str],
    current_category: str | None,
    current_type: str | None,
    expected: tuple[str, str],
):
    assert get_table_type(headers, current_category, current_type) == expected


# 测试车辆信息处理
@pytest.mark.parametrize(
    "car_info,batch_number,expected",
    [
        (
            {
                "产品型号": "TEST001",
                "通用名称": "测试品牌",
                "企业名称": "测试企业",
            },
            "65",
            {
                "型号": "TEST001",
                "品牌": "测试品牌",
                "企业名称": "测试企业",
                "batch": "65",
            },
        ),
        (
            {
                "车辆型号": "TEST002",
                "商标": "测试品牌2",
                "企业名称": "测试企业2",
            },
            "66",
            {
                "型号": "TEST002",
                "品牌": "测试品牌2",
                "企业名称": "测试企业2",
                "batch": "66",
            },
        ),
    ],
)
def test_process_car_info(car_info: dict, batch_number: str, expected: dict):
    result = process_car_info(car_info, batch_number)
    for key, value in expected.items():
        assert result[key] == value


# 测试文档内容提取
def test_extract_doc_content():
    # 创建模拟的Document对象
    mock_doc = Mock()
    mock_paragraphs = [
        Mock(text="第六十五批"),
        Mock(text="一、节能型汽车"),
        Mock(text="（一）乘用车"),
        Mock(text="关于某些说明"),
        Mock(text="\n\n"),  # 添加两个空行来分隔不同的额外信息
        Mock(text="技术要求：其他内容"),  # 添加一个新的额外信息
        Mock(text="\n\n"),  # 添加两个空行来分隔不同的额外信息
    ]
    mock_doc.paragraphs = mock_paragraphs

    # 模拟 element.body
    mock_para_elements = []
    for para in mock_paragraphs:
        mock_element = Mock()
        mock_element.tag = Mock()
        mock_element.tag.endswith = Mock(side_effect=lambda x: x == "p")
        mock_element.text = para.text
        mock_para_elements.append(mock_element)

    mock_doc.element = Mock()
    mock_doc.element.body = mock_para_elements

    with patch("main.Document", return_value=mock_doc):
        paragraphs, extra_info = extract_doc_content("test.docx")

        # 验证段落内容
        assert "第六十五批" in paragraphs
        assert "一、节能型汽车" in paragraphs
        assert "（一）乘用车" in paragraphs

        # 验证额外信息
        assert len(extra_info) == 2
        assert any(
            info["type"] == "政策" and "关于某些说明" in info["content"]
            for info in extra_info
        )
        assert any(
            info["type"] == "说明" and "技术要求" in info["content"]
            for info in extra_info
        )

        # 验证额外信息的内容
        for info in extra_info:
            if info["type"] == "政策":
                assert "关于某些说明" in info["content"]
                assert "技术要求" not in info["content"]
            elif info["type"] == "说明":
                assert "技术要求" in info["content"]
                assert "关于某些说明" not in info["content"]


# 测试DocProcessor类
def test_doc_processor():
    # 创建模拟的Document对象
    mock_doc = Mock()

    # 创建表格
    mock_table = Mock(spec=Table)
    mock_cells1 = [
        Mock(text="企业名称"),
        Mock(text="型号"),
        Mock(text="排量(ml)"),
        Mock(text="综合燃料消耗量"),
    ]
    mock_cells2 = [
        Mock(text="测试企业"),
        Mock(text="TEST001"),
        Mock(text="1998"),
        Mock(text="6.5"),
    ]
    mock_row1 = Mock()
    mock_row1.cells = mock_cells1
    mock_row2 = Mock()
    mock_row2.cells = mock_cells2
    mock_table.rows = [mock_row1, mock_row2]
    mock_doc.tables = [mock_table]

    # 创建段落
    mock_paragraphs = [
        Mock(text="第六十五批"),
        Mock(text="一、节能型汽车"),
        Mock(text="（一）乘用车"),
    ]
    mock_doc.paragraphs = mock_paragraphs

    # 模拟 element.body
    mock_table_element = Mock()
    mock_table_element.tag = Mock()
    mock_table_element.tag.endswith = Mock(side_effect=lambda x: x == "tbl")
    mock_table_element._element = mock_table  # 使用 _element 作为表格的引用
    mock_table._element = mock_table_element  # 设置表格的 _element 属性

    mock_para_elements = []
    for para in mock_paragraphs:
        mock_element = Mock()
        mock_element.tag = Mock()
        mock_element.tag.endswith = Mock(side_effect=lambda x: x == "p")
        mock_element.text = para.text
        mock_para_elements.append(mock_element)

    mock_doc.element = Mock()
    mock_doc.element.body = mock_para_elements + [mock_table_element]

    # 模拟表格的 _element 属性
    mock_table._element = mock_table_element
    mock_table_element._element = mock_table

    with patch("main.Document", return_value=mock_doc):
        processor = DocProcessor("test.docx")
        cars = processor.process()

        assert len(cars) > 0
        assert cars[0]["car_type"] == 2  # 节能型
        assert "batch" in cars[0]
        assert cars[0]["batch"] == "65"
        assert cars[0]["企业名称"] == "测试企业"
        assert cars[0]["型号"] == "TEST001"


# 测试完整的处理流程
def test_process_command(tmp_path: Path):
    # 创建测试文件和目录
    test_dir = tmp_path / "test_docs"
    test_dir.mkdir()
    test_file = test_dir / "test.docx"
    test_file.touch()

    # 创建模拟的Document对象
    mock_doc = Mock()

    # 创建表格
    mock_table = Mock(spec=Table)
    mock_cells1 = [
        Mock(text="企业名称"),
        Mock(text="型号"),
        Mock(text="排量(ml)"),
        Mock(text="综合燃料消耗量"),
    ]
    mock_cells2 = [
        Mock(text="测试企业"),
        Mock(text="TEST001"),
        Mock(text="1998"),
        Mock(text="6.5"),
    ]
    mock_row1 = Mock()
    mock_row1.cells = mock_cells1
    mock_row2 = Mock()
    mock_row2.cells = mock_cells2
    mock_table.rows = [mock_row1, mock_row2]
    mock_doc.tables = [mock_table]

    # 创建段落
    mock_paragraphs = [
        Mock(text="第六十五批"),
        Mock(text="一、节能型汽车"),
        Mock(text="（一）乘用车"),
    ]
    mock_doc.paragraphs = mock_paragraphs

    # 模拟 element.body
    mock_table_element = Mock()
    mock_table_element.tag = Mock()
    mock_table_element.tag.endswith = Mock(side_effect=lambda x: x == "tbl")
    mock_table_element._element = mock_table  # 使用 _element 作为表格的引用
    mock_table._element = mock_table_element  # 设置表格的 _element 属性

    mock_para_elements = []
    for para in mock_paragraphs:
        mock_element = Mock()
        mock_element.tag = Mock()
        mock_element.tag.endswith = Mock(side_effect=lambda x: x == "p")
        mock_element.text = para.text
        mock_para_elements.append(mock_element)

    mock_doc.element = Mock()
    mock_doc.element.body = mock_para_elements + [mock_table_element]

    # 模拟表格的 _element 属性
    mock_table._element = mock_table_element
    mock_table_element._element = mock_table

    # 创建一个模拟的DataFrame
    mock_df = Mock()
    mock_df.columns = ["batch", "car_type", "型号", "企业名称", "raw_text"]
    mock_df.__getitem__ = Mock(return_value=mock_df)
    mock_df.to_csv = Mock()
    mock_df.__len__ = Mock(return_value=1)
    mock_df.__iter__ = Mock(return_value=iter([{"batch": "65", "car_type": 2}]))

    with (
        patch("main.Document", return_value=mock_doc),
        patch("pandas.DataFrame", return_value=mock_df),
        patch("main.console.print") as mock_print,
    ):
        from main import process_files

        # 直接调用处理函数
        process_files(
            str(test_dir), "output.csv", verbose=True, preview=False, compare=None
        )

        # 验证是否调用了console.print
        mock_print.assert_called()

        # 验证是否创建了DataFrame并保存了文件
        mock_df.to_csv.assert_called_once_with(
            "output.csv", index=False, encoding="utf-8-sig"
        )


if __name__ == "__main__":
    pytest.main(["-v", "--cov=main", "--cov-report=term-missing"])
