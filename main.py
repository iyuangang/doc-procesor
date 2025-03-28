from pathlib import Path
import pandas as pd  # type: ignore
from docx import Document  # type: ignore
from docx.document import Document as DocxDocument
import re
from typing import Dict, Any, Optional, List, Set, Tuple, Callable, Union
import click
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.progress import (
    Progress,
    SpinnerColumn,
    BarColumn,
    TimeRemainingColumn,
    TimeElapsedColumn,
)
from rich.text import Text
from rich.tree import Tree
import textwrap
from functools import lru_cache, partial
import cProfile
import pstats
from io import StringIO
import time
import psutil
import os
import gc
import logging
import multiprocessing as mp
import tempfile
import shutil
import yaml
import logging.config
from datetime import datetime
from dataclasses import dataclass, field


# 创建控制台对象
console = Console()


# 预编译正则表达式
BATCH_NUMBER_PATTERN = re.compile(r"第([一二三四五六七八九十百零\d]+)批")
WHITESPACE_PATTERN = re.compile(r"\s+")
CHINESE_NUMBER_PATTERN = re.compile(r"([一二三四五六七八九十百零]+)")

# 中文数字映射表
CN_NUMS = {
    "零": "0",
    "一": "1",
    "二": "2",
    "三": "3",
    "四": "4",
    "五": "5",
    "六": "6",
    "七": "7",
    "八": "8",
    "九": "9",
    "十": "10",
    "百": "100",
}


@lru_cache(maxsize=1024)
def cn_to_arabic(cn_num: str) -> str:
    """
    将中文数字转换为阿拉伯数字, 使用缓存提高性能
    """
    if cn_num.isdigit():
        return cn_num

    # 处理个位数
    if len(cn_num) == 1:
        return CN_NUMS.get(cn_num, cn_num)

    # 处理"百"开头的数字
    if "百" in cn_num:
        parts = cn_num.split("百")
        hundreds = int(CN_NUMS[parts[0]])
        if not parts[1]:  # 整百
            return str(hundreds * 100)
        # 处理带"零"的情况
        if parts[1].startswith("零"):
            ones = int(CN_NUMS[parts[1][-1]])
            return str(hundreds * 100 + ones)
        return str(hundreds * 100 + int(cn_to_arabic(parts[1])))

    # 处理"十"开头的数字
    if cn_num.startswith("十"):
        if len(cn_num) == 1:
            return "10"
        return "1" + CN_NUMS[cn_num[1]]

    # 处理带十的两位数
    if "十" in cn_num:
        parts = cn_num.split("十")
        # 直接在字符串格式化中使用CN_NUMS字典的值, 避免类型转换
        if len(parts) == 1 or not parts[1]:
            return f"{CN_NUMS[parts[0]]}0"
        return f"{CN_NUMS[parts[0]]}{CN_NUMS[parts[1]]}"

    return CN_NUMS.get(cn_num, cn_num)


@lru_cache(maxsize=1024)
def extract_batch_number(text: str) -> Optional[str]:
    """
    从文本中提取批次号, 使用缓存提高性能
    """
    # 先尝试匹配完整的批次号格式
    match = BATCH_NUMBER_PATTERN.search(text)
    if match:
        num = match.group(1)
        # 如果是纯数字, 直接返回
        if num.isdigit():
            return num

        # 转换中文数字
        try:
            return cn_to_arabic(num)
        except (KeyError, ValueError):
            return None

    # 如果没有找到批次号格式, 尝试直接转换纯中文数字
    if any(char in text for char in "一二三四五六七八九十百零"):
        try:
            # 提取连续的中文数字
            match = CHINESE_NUMBER_PATTERN.search(text)
            if match:
                return cn_to_arabic(match.group(1))
        except (KeyError, ValueError):
            pass

    return None


@lru_cache(maxsize=1024)
def clean_text(text: str) -> str:
    """
    清理文本内容, 使用缓存提高性能
    """
    # 移除多余的空白字符
    text = WHITESPACE_PATTERN.sub(" ", text.strip())
    # 统一全角字符到半角
    text = text.replace(", ", ",").replace("；", ";")
    return text


def validate_car_info(
    car_info: Dict[str, Any],
) -> Tuple[bool, str, Optional[Dict[str, Any]]]:
    """验证并尝试修复车辆信息"""
    # 基本验证
    if not car_info or not any(str(value).strip() for value in car_info.values()):
        return False, "空行", None

    # 检查是否为合计行
    if any(
        str(value).strip().startswith(("合计", "总计")) for value in car_info.values()
    ):
        return False, "合计行", None

    # 尝试修复数据
    fixed_info = car_info.copy()

    # 1. 处理变速器信息
    if "型式" in fixed_info and "档位数" in fixed_info:
        fixed_info["变速器"] = f"{fixed_info.pop('型式')} {fixed_info.pop('档位数')}"

    # 2. 标准化数值字段
    numeric_fields = ["排量(ml)", "整车整备质量(kg)", "综合燃料消耗量（L/100km）"]
    for fields in numeric_fields:
        if fields in fixed_info:
            value = fixed_info[fields]
            if isinstance(value, str):
                # 处理多个数值的情况（如范围值）
                if "/" in value:
                    values = [float(v.strip()) for v in value.split("/") if v.strip()]
                    fixed_info[fields] = min(values)  # 使用最小值
                else:
                    try:
                        fixed_info[fields] = float(value.replace(", ", ","))
                    except ValueError:
                        logging.warning(f"无法转换数值: {fields}={value}")

    # 3. 确保必要字段存在
    required_fields = ["car_type", "category", "sub_type"]
    for fields in required_fields:
        if fields not in fixed_info:
            return False, f"缺少必要字段: {fields}", None

    return True, "", fixed_info


def get_table_type(
    headers: List[str], current_category: Optional[str], current_type: Optional[str]
) -> Tuple[str, str]:
    """
    根据表头判断表格类型，使用当前上下文确定子类型

    Args:
        headers: 表头列表
        current_category: 当前分类（节能型或新能源）
        current_type: 当前子类型（表格所在的子分类）

    Returns:
        (category, sub_type)元组
    """
    # 标准化表头
    normalized_headers = [h.strip().lower() for h in headers]

    # 验证必要的列是否存在
    required_columns = {"序号", "企业名称"}
    missing_columns = required_columns - set(normalized_headers)
    if missing_columns:
        raise ValueError(f"表格缺少必要的列: {missing_columns}")

    # 处理特殊的表头组合
    if "型式" in normalized_headers and "档位数" in normalized_headers:
        # 合并为变速器列
        idx = normalized_headers.index("型式")
        normalized_headers[idx] = "变速器"
        normalized_headers.pop(idx + 1)

    header_text = " ".join(normalized_headers).lower()

    # 只从表头判断category（节能型或新能源）
    category = current_category or "未知"

    # 节能型车辆的特征关键词
    energy_saving_indicators = [
        "排量(ml)",
        "燃料消耗量",
        "排量",
        "油耗",
        "发动机",
        "cng",
        "lng",
        "燃气",
        "天然气",
        "燃料种类",
    ]

    # 新能源车辆的特征关键词
    new_energy_indicators = [
        "电池",
        "电动",
        "纯电动续驶里程",
        "动力电池",
        "电量",
        "电机",
        "混合动力",
        "燃料电池",
        "充电",
    ]

    # 检查表头是否包含节能型车辆特征
    for indicator in energy_saving_indicators:
        if indicator in header_text:
            category = "节能型"
            break

    # 检查表头是否包含新能源车辆特征
    for indicator in new_energy_indicators:
        if indicator in header_text:
            category = "新能源"
            break

    # 如果在明确的节能型部分中，优先使用节能型分类
    if "节能型" in str(current_category).lower():
        category = "节能型"

    # 如果在明确的新能源部分中，优先使用新能源分类
    if "新能源" in str(current_category).lower():
        category = "新能源"

    # 始终使用当前上下文的子类型，不从表头判断
    sub_type = current_type or "未知"

    # 确保返回的是字符串类型
    return str(category), str(sub_type)


def process_car_info(
    car_info: Dict[str, Any], batch_number: Optional[str] = None
) -> Dict[str, Any]:
    """
    处理车辆信息, 合并和标准化字段

    Args:
        car_info: 原始车辆信息字典
        batch_number: 批次号

    Returns:
        处理后的车辆信息字典
    """
    # 添加批次号
    if batch_number:
        car_info["batch"] = batch_number

    # 合并型号字段
    model_fields = ["产品型号", "车辆型号", "型号"]
    model_values = []
    for fields in model_fields:
        if fields in car_info:
            value = car_info.pop(fields) if fields != "型号" else car_info.get(fields)
            if value and str(value).strip():
                model_values.append(clean_text(str(value)))

    if model_values:
        car_info["型号"] = model_values[0]  # 使用第一个非空的型号

    # 标准化字段名称
    field_mapping: Dict[str, str] = {
        "通用名称": "品牌",
        "商标": "品牌",
        "生产企业": "企业名称",
        "企业": "企业名称",
    }

    # 处理字段映射
    for old_field, new_field in field_mapping.items():
        if old_field in car_info:
            value = car_info.pop(old_field)
            if value and str(value).strip():
                car_info[new_field] = clean_text(str(value))

    # 清理其他字段的文本, 但保留所有值
    for key in car_info:
        if isinstance(car_info[key], str):
            car_info[key] = clean_text(car_info[key])

    return car_info


def extract_doc_content(doc_path: str) -> Tuple[List[str], List[Dict[str, str]]]:
    """
    提取文档中除表格外的内容, 并分离额外信息
    """
    doc: DocxDocument = Document(doc_path)
    paragraphs: List[str] = []
    extra_info: List[Dict[str, str]] = []
    current_section: Optional[str] = None
    batch_found = False
    batch_number = None

    # 额外信息的标识词和对应类型
    info_types: Dict[str, str] = {
        "勘误": "勘误",
        "关于": "政策",
        "符合": "说明",
        "技术要求": "说明",
        "自动转入": "说明",
        "第二部分": "说明",
    }

    # 用于收集连续的额外信息文本
    current_extra_info: Optional[Dict[str, str]] = None

    def save_current_extra_info() -> None:
        """保存当前的额外信息"""
        nonlocal current_extra_info
        if current_extra_info:
            # 清理和规范化内容
            content = current_extra_info["content"]
            # 移除多余的空白字符
            content = re.sub(r"\s+", " ", content)
            # 移除换行符
            content = content.replace("\n", " ")
            current_extra_info["content"] = content.strip()

            # 添加批次号
            if batch_number:
                current_extra_info["batch"] = batch_number

            # 检查是否需要合并相同类型和章节的信息
            for info in extra_info:
                if (
                    info["type"] == current_extra_info["type"]
                    and info["section"] == current_extra_info["section"]
                ):
                    info["content"] = (
                        info["content"] + " " + current_extra_info["content"]
                    )
                    current_extra_info = None
                    return

            extra_info.append(current_extra_info)
            current_extra_info = None

    # 遍历文档段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # 如果遇到空行, 保存当前的额外信息
            if current_extra_info:
                save_current_extra_info()
            continue

        # 检查批次号
        if not batch_found and "批" in text:
            extracted_batch = extract_batch_number(text)
            if extracted_batch:
                batch_number = extracted_batch
                paragraphs.append(text)  # 将批次号信息放在最前面
                batch_found = True
                continue

        # 识别主要分类
        if "节能型汽车" in text or "新能源汽车" in text:
            save_current_extra_info()
            current_section = text
            paragraphs.append(text)
        # 识别子分类,排除括号中有数字的
        elif text.startswith("（") and not any(str.isdigit() for str in text):
            save_current_extra_info()
            current_section = text
            paragraphs.append(text)
        # 识别额外信息
        elif any(marker in text for marker in info_types.keys()):
            # 如果当前文本包含新的标识词, 保存之前的信息并创建新的
            if current_extra_info:
                save_current_extra_info()

            # 创建新的额外信息
            info_type = next((t for m, t in info_types.items() if m in text), "其他")
            current_extra_info = {
                "section": current_section or "文档说明",
                "type": info_type,
                "content": text,
            }
        # 如果当前有未处理的额外信息, 将文本追加到内容中
        elif current_extra_info is not None:
            current_extra_info["content"] = current_extra_info["content"] + " " + text
        else:
            paragraphs.append(text)

    # 保存最后一条未处理的额外信息
    save_current_extra_info()

    return paragraphs, extra_info


def print_docx_content(doc_path: str) -> None:
    """打印文档内容预览, 显示所有元素的详细信息"""
    try:
        doc: DocxDocument = Document(doc_path)
        console.print(
            Panel(
                f"[bold cyan]文件详细内容: {doc_path}[/bold cyan]", border_style="cyan"
            )
        )

        # 创建一个树形结构
        tree = Tree(f"📄 {Path(doc_path).name}", style="bold blue")

        # 添加段落内容
        paragraphs_node = tree.add("[bold magenta]📝 段落内容[/bold magenta]")
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else "默认样式"
                para_node = paragraphs_node.add(
                    f"[blue]段落 {i}[/blue] ([yellow]{style_name}[/yellow])"
                )
                if "批" in text:
                    para_node.add(f"🔖 [bold red]{text}[/bold red]")
                elif "节能型汽车" in text or "新能源汽车" in text:
                    para_node.add(f"📌 [bold green]{text}[/bold green]")
                elif text.startswith("（") and not any(str.isdigit() for str in text):
                    para_node.add(f"📎 [bold yellow]{text}[/bold yellow]")
                elif any(
                    marker in text
                    for marker in ["勘误", "关于", "符合", "技术要求", "自动转入"]
                ):
                    para_node.add(f"ℹ️ [bold magenta]{text}[/bold magenta]")
                else:
                    para_node.add(Text(textwrap.shorten(text, width=100)))

        # 添加表格内容
        tables_node = tree.add("[bold cyan]📊 表格内容[/bold cyan]")
        for i, table in enumerate(doc.tables, 1):
            if table.rows:
                table_node = tables_node.add(
                    f"[blue]表格 {i}[/blue] ({len(table.rows)}行 x {len(table.rows[0].cells)}列)"
                )

                # 创建表格预览
                preview_table = Table(
                    show_header=True,
                    header_style="bold green",
                    border_style="blue",
                    title=f"表格 {i} 预览",
                    title_style="bold cyan",
                )

                # 添加表头
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                for header in headers:
                    preview_table.add_column(header, overflow="fold")

                # 添加数据行预览
                for row_idx, row in enumerate(table.rows[1:6], 1):  # 只显示前5行数据
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):  # 跳过空行
                        preview_table.add_row(*cells)

                table_node.add(preview_table)

        console.print(tree)

    except Exception as e:
        console.print(
            Panel(
                f"[bold red]预览文件 {doc_path} 时出错: {e}[/bold red]",
                border_style="red",
            )
        )


def display_statistics(
    total_count: int, energy_saving_count: int, new_energy_count: int, output_file: str
) -> None:
    """Display processing statistics in a formatted table."""
    # 在显示表格前添加标题, 表明这是关键信息
    console.print()
    console.print("[bold cyan]📊 关键信息：处理统计报告[/bold cyan]")

    # 创建统计表格
    stats_table = Table(
        title="📊 处理统计报告",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    stats_table.add_column("项目", style="cyan")
    stats_table.add_column("数值", justify="right", style="green")
    stats_table.add_column("占比", justify="right", style="yellow")

    # 计算百分比
    energy_saving_percent = (
        energy_saving_count / total_count * 100 if total_count > 0 else 0
    )
    new_energy_percent = new_energy_count / total_count * 100 if total_count > 0 else 0

    # 添加行
    stats_table.add_row("📝 总记录数", f"{total_count:,}", "100%")
    stats_table.add_row(
        "🚗 节能型汽车", f"{energy_saving_count:,}", f"{energy_saving_percent:.1f}%"
    )
    stats_table.add_row(
        "⚡ 新能源汽车", f"{new_energy_count:,}", f"{new_energy_percent:.1f}%"
    )
    stats_table.add_row("💾 输出文件", output_file, "")

    # 显示表格
    console.print(stats_table)
    console.print()


@dataclass
class DocumentNode:
    """文档节点类, 用于构建文档树结构"""

    title: str
    level: int
    node_type: str  # 'section', 'subsection', 'table', 'text', 'note', 'correction'
    content: Optional[str] = None
    batch_number: Optional[str] = None
    children: List["DocumentNode"] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocumentStructure:
    """文档结构类, 用于构建和管理文档的层级结构"""

    def __init__(self) -> None:
        self.root = DocumentNode("文档结构", 0, "root")
        self.current_section: Optional[DocumentNode] = None
        self.current_subsection: Optional[DocumentNode] = None
        self.batch_number: Optional[str] = None

    def add_node(
        self,
        title: str,
        node_type: str,
        content: Optional[str] = None,
        level: Optional[int] = None,
        metadata: Optional[Dict[str, Any]] = None,
        parent_node: Optional[DocumentNode] = None,
    ) -> DocumentNode:
        """添加新节点到文档树"""
        if level is None:
            if node_type == "section":
                level = 1
            elif node_type == "subsection":
                level = 2
            elif node_type == "numbered_section":
                level = 3
            elif node_type == "numbered_subsection":
                level = 4
            else:
                level = 5

        node = DocumentNode(
            title=title,
            level=level,
            node_type=node_type,
            content=content,
            batch_number=self.batch_number,
            metadata=metadata or {},
        )

        # 如果指定了父节点, 直接添加到父节点
        if parent_node:
            parent_node.children.append(node)
            return node

        # 否则使用默认的层级逻辑
        if level == 1:
            self.root.children.append(node)
        elif level == 2:
            if self.current_section:
                self.current_section.children.append(node)
            else:
                self.root.children.append(node)
        else:
            if self.current_subsection:
                self.current_subsection.children.append(node)
            elif self.current_section:
                self.current_section.children.append(node)
            else:
                self.root.children.append(node)

        return node

    def set_batch_number(self, batch_number: str) -> None:
        """设置批次号"""
        self.batch_number = batch_number

    def to_dict(self) -> Dict[str, Any]:
        """将文档结构转换为字典格式"""

        def node_to_dict(node: DocumentNode) -> Dict[str, Any]:
            return {
                "title": node.title,
                "type": node.node_type,
                "level": node.level,
                "content": node.content,
                "batch_number": node.batch_number,
                "metadata": node.metadata,
                "children": [node_to_dict(child) for child in node.children],
            }

        return node_to_dict(self.root)


def display_doc_content(doc_structure: DocumentStructure) -> None:
    """使用树形结构显示文档内容"""

    def get_node_style(node: DocumentNode) -> Tuple[str, str]:
        """获取节点的样式和图标"""
        styles = {
            "root": ("bold blue", "📑"),
            "section": ("bold cyan", "📌"),
            "subsection": ("bold yellow", "📎"),
            "numbered_section": ("bold green", "🔢"),
            "numbered_subsection": ("bold magenta", "📍"),
            "table": ("bold blue", "📊"),
            "text": ("white", "📝"),
            "note": ("bold magenta", "ℹ️"),
            "correction": ("bold red", "⚠️"),
        }
        return styles.get(node.node_type, ("white", "•"))

    def add_node_to_tree(tree: Tree, node: DocumentNode) -> None:
        """递归添加节点到树中"""
        style, icon = get_node_style(node)

        # 构建节点标题
        title = f"{icon} {node.title}"
        if node.batch_number and node.level <= 1:
            title += f" [dim](第{node.batch_number}批)[/dim]"

        # 创建节点
        branch = tree.add(f"[{style}]{title}[/{style}]")

        # 添加内容（如果有且与标题不同）
        if node.content and node.content != node.title:
            content_lines = textwrap.wrap(node.content, width=100)
            for line in content_lines:
                branch.add(f"[dim]{line}[/dim]")

        # 添加元数据（如果有）
        if node.metadata:
            meta_branch = branch.add("[dim]元数据[/dim]")
            for key, value in node.metadata.items():
                meta_branch.add(f"[dim]{key}: {value}[/dim]")

        # 递归处理子节点
        for child in node.children:
            add_node_to_tree(branch, child)

    # 创建主树
    tree = Tree("📄 文档结构", style="bold blue")
    for child in doc_structure.root.children:
        add_node_to_tree(tree, child)

    # 显示树
    console.print("\n")
    console.print(Panel(tree, border_style="blue"))
    console.print()


def display_comparison(new_models: Set[str], removed_models: Set[str]) -> None:
    """显示型号对比结果"""
    # 创建对比表格
    compare_table = Table(
        title="🔄 型号对比",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    compare_table.add_column("变更类型", style="cyan")
    compare_table.add_column("数量", justify="right", style="green")
    compare_table.add_column("型号列表", style="yellow")

    # 添加新增型号
    if new_models:
        models_text = "\n".join(f"✨ {model}" for model in sorted(new_models))
        compare_table.add_row("➕ 新增", str(len(new_models)), models_text)

    # 添加移除型号
    if removed_models:
        models_text = "\n".join(f"❌ {model}" for model in sorted(removed_models))
        compare_table.add_row("➖ 移除", str(len(removed_models)), models_text)

    if new_models or removed_models:
        console.print()
        console.print(compare_table)
        console.print()
    else:
        console.print(Panel("[green]✅ 没有型号变更[/green]", border_style="green"))


@click.group()
def cli() -> None:
    """处理车辆数据文档的命令行工具"""
    pass


def extract_car_info(doc_path: str, verbose: bool = False) -> List[Dict[str, Any]]:
    """从docx文件中提取车辆信息"""
    processor = DocProcessor(doc_path, verbose)
    result: List[Dict[str, Any]] = processor.process()
    return result


def get_memory_usage() -> str:
    """获取当前进程的内存使用情况"""
    process = psutil.Process(os.getpid())
    memory_info = process.memory_info()
    return f"{memory_info.rss / 1024 / 1024:.1f}MB"


def process_doc(
    doc_path: str, verbose: bool = False, config: Optional[dict] = None
) -> List[Dict[str, Any]]:
    """单个文档处理函数, 用于多进程"""
    try:
        # 如果是简洁模式, 抑制详细输出但保留三项关键信息
        processor = DocProcessor(doc_path, verbose, config)
        result: List[Dict[str, Any]] = processor.process()
        return result
    except Exception as e:
        logging.error(f"处理文档 {doc_path} 失败: {str(e)}")
        return []


def verify_all_batches(all_cars_data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """验证所有批次的数据一致性"""
    # 按批次分组
    batch_data: Dict[str, List[Dict[str, Any]]] = {}
    for car in all_cars_data:
        batch = car.get("batch")
        if not batch:
            continue

        if batch not in batch_data:
            batch_data[batch] = []
        batch_data[batch].append(car)

    # 验证每个批次
    results = {}
    for batch, cars in batch_data.items():
        # 按表格分组
        table_counts = {}
        for car in cars:
            table_id = car.get("table_id", "未知")
            if table_id not in table_counts:
                table_counts[table_id] = 0
            table_counts[table_id] += 1

        # 总计
        total_count = len(cars)

        results[batch] = {"total": total_count, "table_counts": table_counts}

    return results


def display_batch_verification(batch_results: Dict[str, Any]) -> None:
    """显示批次验证结果"""
    if not batch_results:
        console.print(
            Panel(
                "[yellow]⚠️ 没有批次数据可供验证[/yellow]",
                title="批次验证",
                border_style="yellow",
            )
        )
        return

    # 创建批次汇总表格
    summary_table = Table(
        title="🔍 批次数据汇总",
        title_style="bold cyan",
        show_header=True,
        header_style="bold green",
        border_style="blue",
    )

    # 添加列
    summary_table.add_column("批次", style="cyan")
    summary_table.add_column("记录数", justify="right", style="green")
    summary_table.add_column("表格数", justify="right", style="yellow")

    # 计算批次总数, 如果超过一定数量, 只显示部分
    batch_count = len(batch_results)
    show_all = batch_count <= 50  # 只有50个批次以内才全部显示

    # 添加批次数据
    total_records = 0
    total_tables = 0

    sorted_batches = sorted(batch_results.items())

    # 如果批次太多, 只显示前20个和后20个
    if not show_all:
        display_batches = sorted_batches[:20] + sorted_batches[-20:]
        console.print(
            f"[yellow]注意：只显示前20个和后20个批次（共{batch_count}个批次）[/yellow]"
        )
    else:
        display_batches = sorted_batches

    for batch, data in display_batches:
        total_records += data["total"]
        table_count = len(data["table_counts"])
        total_tables += table_count
        summary_table.add_row(f"第{batch}批", str(data["total"]), str(table_count))

    # 如果有省略的批次, 添加省略提示行
    if not show_all and batch_count > 40:
        summary_table.add_row(f"... (省略 {batch_count - 40} 个批次) ...", "...", "...")

    # 添加合计行
    summary_table.add_row(
        "[bold]合计[/bold]",
        f"[bold]{total_records}[/bold]",
        f"[bold]{total_tables}[/bold]",
    )

    # 在表格前添加标题, 表明这是关键信息
    console.print()
    console.print("[bold cyan]📊 关键信息：批次数据汇总[/bold cyan]")
    console.print(summary_table)
    console.print()


def process_files(
    input_path: str,
    output: str,
    verbose: bool = False,
    preview: bool = False,
    compare: Optional[str] = None,
    config: Optional[dict] = None,
) -> None:
    """处理指定的docx文件或目录下的所有docx文件的核心逻辑"""
    logger = logging.getLogger(__name__)

    try:
        input_path_obj = Path(input_path)
        if input_path_obj.is_file():
            if input_path_obj.suffix.lower() != ".docx":
                raise ValueError("指定的文件不是docx文件")
            doc_files = [input_path_obj]
        else:
            doc_files = list(input_path_obj.glob("*.docx"))

        if not doc_files:
            raise ValueError("未找到.docx文件")

        if preview:
            for doc_file in doc_files:
                print_docx_content(str(doc_file))

        # 使用多进程处理文档
        num_processes = min(mp.cpu_count(), len(doc_files))
        logger.info(f"使用 {num_processes} 个进程处理 {len(doc_files)} 个文件")

        with mp.Pool(num_processes) as pool:
            with Progress(
                "[progress.description]{task.description}",
                SpinnerColumn(),
                BarColumn(bar_width=None),
                "[progress.percentage]{task.percentage:>3.1f}%",
                "•",
                "[bold blue]{task.completed}/{task.total}",
                "•",
                TimeRemainingColumn(),
                "•",
                TimeElapsedColumn(),
                console=console,
                transient=True,
            ) as progress:
                main_task = progress.add_task(
                    "[bold cyan]🔄 处理文件", total=len(doc_files)
                )

                # 使用partial固定参数
                process_func = partial(process_doc, verbose=verbose, config=config)

                # 使用imap处理结果
                all_cars_data = []
                error_files = []

                for doc_file, cars in zip(
                    doc_files, pool.imap(process_func, [str(f) for f in doc_files])
                ):
                    if cars:
                        all_cars_data.extend(cars)
                        logger.info(
                            f"✅ 文件 {doc_file} 处理完成, 提取到 {len(cars)} 条记录"
                        )
                    else:
                        error_files.append(doc_file)
                        logger.error(f"❌ 文件 {doc_file} 处理失败")

                    progress.advance(main_task)

                    # 定期清理内存
                    if len(all_cars_data) > 10000:
                        gc.collect()

        # 处理结果
        if all_cars_data:
            try:
                # 验证所有批次的数据一致性
                batch_results = verify_all_batches(all_cars_data)

                # 始终显示批次验证结果, 即使在简洁模式下
                display_batch_verification(batch_results)

                # 估计数据大小
                estimated_size = len(all_cars_data) * 500  # 假设每条记录约500字节
                is_large_dataset = estimated_size > 100 * 1024 * 1024  # 100MB

                if is_large_dataset:
                    logger.info(
                        f"大数据集 ({len(all_cars_data)} 条记录), 使用优化处理..."
                    )

                    # 使用分块处理
                    chunk_size = 50000
                    with open(output, "w", encoding="utf-8-sig") as f:
                        # 写入表头
                        first_batch: List[Dict[str, Any]] = all_cars_data[
                            :100
                        ]  # 取前100条确定字段
                        all_fields: Set[str] = set()
                        for car in first_batch:
                            all_fields.update(car.keys())

                        base_columns = [
                            "batch",
                            "car_type",
                            "category",
                            "sub_type",
                            "序号",
                            "企业名称",
                            "品牌",
                            "型号",
                            "table_id",
                            "raw_text",
                        ]

                        header_fields = [
                            col for col in base_columns if col in all_fields
                        ] + [
                            col for col in sorted(all_fields) if col not in base_columns
                        ]

                        f.write(",".join(header_fields) + "\n")

                        # 分块写入数据
                        for i in range(0, len(all_cars_data), chunk_size):
                            chunk = all_cars_data[i : i + chunk_size]
                            chunk_df = pd.DataFrame(chunk)
                            chunk_df = chunk_df.reindex(columns=header_fields)

                            if i == 0:
                                chunk_df.to_csv(
                                    f, index=False, header=False, encoding="utf-8-sig"
                                )
                            else:
                                chunk_df.to_csv(
                                    f,
                                    index=False,
                                    header=False,
                                    encoding="utf-8-sig",
                                    mode="a",
                                )

                            # 释放内存
                            del chunk_df
                            gc.collect()

                    logger.info(f"💾 处理完成, 保存结果到: {output}")
                    logger.info(f"📊 总记录数: {len(all_cars_data)}")

                    # 计算统计数据
                    energy_saving_count = sum(
                        1 for car in all_cars_data if car.get("car_type") == 2
                    )
                    new_energy_count = sum(
                        1 for car in all_cars_data if car.get("car_type") == 1
                    )

                    # 始终显示统计信息, 即使在简洁模式下
                    display_statistics(
                        len(all_cars_data),
                        energy_saving_count,
                        new_energy_count,
                        output,
                    )
                else:
                    # 原有处理逻辑
                    all_cars_df = pd.DataFrame(all_cars_data)

                    # 优化列顺序设置
                    base_columns = [
                        "batch",
                        "car_type",
                        "category",
                        "sub_type",
                        "序号",
                        "企业名称",
                        "品牌",
                        "型号",
                        "table_id",
                        "raw_text",
                    ]
                    all_columns = all_cars_df.columns.tolist()
                    final_columns = [
                        col for col in base_columns if col in all_columns
                    ] + [col for col in all_columns if col not in base_columns]

                    # 重新排列列并保存
                    all_cars_df = all_cars_df[final_columns]
                    all_cars_df.to_csv(output, index=False, encoding="utf-8-sig")

                    logger.info(f"💾 处理完成, 保存结果到: {output}")
                    logger.info(f"📊 总记录数: {len(all_cars_df)}")

                    # 始终显示统计信息, 即使在简洁模式下
                    display_statistics(
                        len(all_cars_df),
                        len(all_cars_df[all_cars_df["car_type"] == 2]),
                        len(all_cars_df[all_cars_df["car_type"] == 1]),
                        output,
                    )

                # 如果有处理失败的文件, 显示警告
                if error_files:
                    error_msg = "❌ 以下文件处理失败:\n" + "\n".join(
                        f"  • {f}" for f in error_files
                    )
                    logger.warning(error_msg)
                    console.print(
                        Panel(
                            f"[bold yellow]{error_msg}[/bold yellow]",
                            title="⚠️ 警告",
                            border_style="yellow",
                        )
                    )

                # 如果需要对比
                if compare:
                    try:
                        old_df = pd.read_csv(compare, encoding="utf-8-sig")
                        new_models = set(all_cars_df["型号"].unique())
                        old_models = set(old_df["型号"].unique())
                        display_comparison(
                            new_models - old_models, old_models - new_models
                        )
                        logger.info("✅ 完成型号对比")
                    except Exception as e:
                        error_msg = f"对比文件时出错: {str(e)}"
                        logger.error(error_msg)
                        console.print(
                            Panel(
                                f"[bold red]{error_msg}[/bold red]",
                                title="❌ 错误",
                                border_style="red",
                            )
                        )

            except Exception as e:
                error_msg = f"处理结果时出错: {str(e)}"
                logger.error(error_msg)
                console.print(
                    Panel(
                        f"[bold red]{error_msg}[/bold red]",
                        title="❌ 错误",
                        border_style="red",
                    )
                )
        else:
            logger.warning("未找到任何车辆记录")
            console.print(
                Panel(
                    "[bold yellow]未找到任何车辆记录[/bold yellow]",
                    title="⚠️ 警告",
                    border_style="yellow",
                )
            )

    except Exception as e:
        error_msg = f"处理文件时出错: {str(e)}"
        logger.error(error_msg)
        console.print(
            Panel(
                f"[bold red]{error_msg}[/bold red]", title="❌ 错误", border_style="red"
            )
        )


@lru_cache(maxsize=32)
def load_document(doc_path: str) -> DocxDocument:
    """缓存加载的文档对象"""
    return Document(doc_path)


def profile_function(func: Callable[..., Any]) -> Callable[..., Any]:
    def wrapper(*args: Any, **kwargs: Any) -> Any:
        profile = cProfile.Profile()
        try:
            return profile.runcall(func, *args, **kwargs)
        finally:
            s = StringIO()
            stats = pstats.Stats(profile, stream=s).sort_stats("cumulative")
            stats.print_stats(20)  # 显示前20个最耗时的函数调用
            console.print(f"\n[bold cyan]性能分析报告:[/bold cyan]\n{s.getvalue()}")

    return wrapper


def setup_logging(
    default_path: str = "logging.yaml",
    default_level: Union[str, int] = logging.INFO,
    env_key: str = "LOG_CFG",
) -> None:
    """配置日志记录"""
    path = os.getenv(env_key, default_path)
    if os.path.exists(path):
        with open(path, "rt") as f:
            try:
                config = yaml.safe_load(f.read())
                logging.config.dictConfig(config)
            except Exception as e:
                print(f"加载日志配置出错: {e}")
                setup_default_logging(default_level)
    else:
        setup_default_logging(default_level)


def setup_default_logging(level: Any) -> None:
    """设置默认日志配置"""
    log_dir = "logs"
    if not os.path.exists(log_dir):
        os.makedirs(log_dir)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(log_dir, f"doc_processor_{timestamp}.log")

    logging.basicConfig(
        level=level,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )


class ConfigurationError(Exception):
    """配置错误异常"""

    pass


class ProcessingError(Exception):
    """处理错误异常"""

    pass


class DocumentError(Exception):
    """文档错误异常"""

    pass


def load_config(config_path: str = "config.yaml") -> Dict[Any, Any]:
    """加载配置文件"""
    try:
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                config: Dict[Any, Any] = yaml.safe_load(f)
            return config
        return {}
    except Exception as e:
        raise ConfigurationError(f"加载配置文件出错: {str(e)}")


class DocProcessor:
    def __init__(
        self, doc_path: str, verbose: bool = True, config: Optional[dict] = None
    ):
        self.doc_path = doc_path
        self.start_time = time.time()
        self.config = config or {}
        self.logger = logging.getLogger(__name__)
        self.doc_structure = DocumentStructure()

        try:
            self._load_document()
        except Exception as e:
            self.logger.error(f"初始化文档处理器失败: {str(e)}")
            raise DocumentError(f"无法加载文档 {doc_path}: {str(e)}")

        self.current_category: Optional[str] = None
        self.current_type: Optional[str] = None
        self.batch_number: Optional[str] = None
        self._table_cache: Dict[int, List[Dict[str, Any]]] = {}
        self.cars: List[Dict[str, Any]] = []
        self._processing_times: Dict[str, float] = {}
        self.declared_count: Optional[int] = None  # 声明的总记录数

        # 从配置文件加载设置
        self._chunk_size = self.config.get("chunk_size", 1000)
        self.verbose = verbose
        self._cache_size_limit = self.config.get("cache_size_limit", 50 * 1024 * 1024)
        self._cleanup_interval = self.config.get("cleanup_interval", 300)
        # 添加跳过总记录数检查的配置选项
        self._skip_count_check = self.config.get("skip_count_check", False)
        # 设置搜索限制
        self._max_paragraphs_to_search = self.config.get("max_paragraphs_to_search", 30)
        self._max_tables_to_search = self.config.get("max_tables_to_search", 5)

        # 预编译正则表达式
        self._batch_pattern = re.compile(r"第([一二三四五六七八九十百零\d]+)批")
        self._whitespace_pattern = re.compile(r"\s+")
        self._chinese_number_pattern = re.compile(r"([一二三四五六七八九十百零]+)")
        self._count_pattern = re.compile(
            r"(共计|总计|合计).*?(\d+).*?(款|个|种|辆|台|项)"
        )  # 总记录数模式

        self._last_cache_cleanup = time.time()
        self.logger.info(f"初始化文档处理器: {doc_path}")

        self.current_section: Optional[DocumentNode] = None
        self.current_subsection: Optional[DocumentNode] = None
        self.current_numbered_section: Optional[DocumentNode] = (
            None  # 用于跟踪带数字编号的节点
        )

    def _load_document(self) -> None:
        """安全加载文档, 处理大文件"""
        try:
            file_size = os.path.getsize(self.doc_path)
            self.logger.info(
                f"加载文档 {self.doc_path}, 大小: {file_size/1024/1024:.2f}MB"
            )

            if file_size > 100 * 1024 * 1024:  # 100MB
                self.logger.warning("文档大小超过100MB, 使用临时文件处理")
                with tempfile.NamedTemporaryFile(delete=False) as tmp:
                    shutil.copy2(self.doc_path, tmp.name)
                    self.doc = Document(tmp.name)
                    os.unlink(tmp.name)
            else:
                self.doc = Document(self.doc_path)
        except Exception as e:
            self.logger.error(f"加载文档失败: {str(e)}")
            raise DocumentError(f"无法加载文档 {self.doc_path}: {str(e)}")

    def _check_and_cleanup_cache(self) -> None:
        """检查并清理缓存"""
        current_time = time.time()
        if current_time - self._last_cache_cleanup > self._cleanup_interval:
            cache_size = sum(len(str(v)) for v in self._table_cache.values())
            if cache_size > self._cache_size_limit:
                self._table_cache.clear()
                gc.collect()
            self._last_cache_cleanup = current_time

    def _extract_table_cells_fast(self, table: Any) -> List[List[str]]:
        """优化的表格提取方法"""
        try:
            rows = []
            header_processed = False
            last_company = ""
            last_brand = ""

            # 使用lxml的xpath直接提取文本
            for row in table._tbl.xpath(".//w:tr"):
                cells = []
                for cell in row.xpath(".//w:tc"):
                    # 直接获取所有文本节点
                    text = "".join(t.text for t in cell.xpath(".//w:t"))
                    cells.append(text.strip())

                if not header_processed:
                    processed_headers = self._process_merged_headers(cells)
                    rows.append(processed_headers)
                    header_processed = True
                    continue

                processed_row = self._process_data_row(cells, last_company, last_brand)
                if processed_row:
                    if processed_row[1]:
                        last_company = processed_row[1]
                    if processed_row[2]:
                        last_brand = processed_row[2]
                    rows.append(processed_row)

                # 定期检查缓存
                self._check_and_cleanup_cache()

            return rows
        except Exception as e:
            logging.error(f"表格提取错误: {str(e)}")
            return []

    def _process_merged_headers(self, headers: List[str]) -> List[str]:
        """处理合并的表头, 例如将'型式'和'档位数'合并为'变速器'"""
        processed = []
        i = 0
        while i < len(headers):
            if (
                i + 1 < len(headers)
                and headers[i] == "型式"
                and headers[i + 1] == "档位数"
            ):
                processed.append("变速器")
                i += 2
            else:
                processed.append(headers[i])
                i += 1
        return processed

    def _process_data_row(
        self, row: List[str], last_company: str, last_brand: str
    ) -> Optional[List[str]]:
        """处理数据行, 包括空值处理和数据继承"""
        # 跳过全空行
        if not any(cell.strip() for cell in row):
            return None

        # 处理合计行
        if any(cell.strip().startswith(("合计", "总计")) for cell in row):
            return None

        processed = []
        for i, cell in enumerate(row):
            value = cell.strip()
            if i == 1 and not value:  # 企业名称为空
                processed.append(last_company)
            elif i == 2 and not value:  # 品牌/通用名称为空
                processed.append(last_brand)
            else:
                processed.append(value)

        return processed

    def _extract_declared_count(self) -> Optional[int]:
        """
        从文档中提取批次声明的总记录数
        优化版本：限制搜索范围并提供跳过选项

        Returns:
            Optional[int]: 声明的总记录数, 如果未找到则返回None
        """
        # 如果配置了跳过总记录数检查, 直接返回None
        if self._skip_count_check:
            self.logger.info("根据配置跳过总记录数检查")
            return None

        # 获取文件大小, 如果超过阈值直接跳过
        try:
            file_size = os.path.getsize(self.doc_path) / (1024 * 1024)  # MB
            if file_size > 50:  # 超过50MB的文档
                self.logger.info(
                    f"文档大小 {file_size:.2f}MB 超过阈值, 跳过总记录数检查"
                )
                return None
        except Exception as e:
            self.logger.error(f"获取文件大小失败: {str(e)}")
            pass  # 如果无法获取文件大小, 继续检查

        start_time = time.time()

        # 1. 只搜索前N个段落
        paragraphs_to_search = min(
            self._max_paragraphs_to_search, len(self.doc.paragraphs)
        )
        self.logger.debug(f"搜索前 {paragraphs_to_search} 个段落以寻找总记录数")

        for i, para in enumerate(self.doc.paragraphs[:paragraphs_to_search]):
            text = para.text.strip()
            if not text:
                continue

            if "总" in text or "共" in text or "合计" in text:
                match = self._count_pattern.search(text)
                if match:
                    try:
                        count = int(match.group(2))
                        search_time = time.time() - start_time
                        self.logger.info(
                            f"从段落中提取到总记录数: {count} (搜索耗时: {search_time:.2f}秒)"
                        )
                        return count
                    except (ValueError, IndexError):
                        continue

        # 2. 只搜索前M个表格
        tables_to_search = min(self._max_tables_to_search, len(self.doc.tables))
        self.logger.debug(f"搜索前 {tables_to_search} 个表格以寻找总记录数")

        for i, table in enumerate(self.doc.tables[:tables_to_search]):
            if not table.rows:
                continue

            # 只检查表格的前3行和后3行, 这些位置最可能出现合计信息
            rows_to_check = []
            if len(table.rows) > 6:
                rows_to_check = list(table.rows[:3]) + list(table.rows[-3:])
            else:
                rows_to_check = list(table.rows)

            for row in rows_to_check:
                cells = [cell.text.strip() for cell in row.cells]
                # 检查是否包含合计相关的内容
                if any(cell.startswith(("合计", "总计")) for cell in cells):
                    # 尝试从合计行中获取数值
                    for cell in cells:
                        if cell.isdigit():
                            count = int(cell)
                            search_time = time.time() - start_time
                            self.logger.info(
                                f"从表格合计行中提取到总记录数: {count} (搜索耗时: {search_time:.2f}秒)"
                            )
                            return count

        search_time = time.time() - start_time
        self.logger.warning(f"未能找到批次总记录数声明 (搜索耗时: {search_time:.2f}秒)")
        return None

    def _extract_car_info(
        self, table_index: int, batch_number: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """从表格中提取车辆信息, 使用优化的处理方式"""
        # 检查缓存
        if table_index in self._table_cache:
            return self._table_cache[table_index]

        start_time = time.time()
        table_cars: List[Dict[str, Any]] = []
        table = self.doc.tables[table_index]

        if not table or not table.rows:
            return table_cars

        # 使用快速方法提取所有单元格内容
        all_rows = self._extract_table_cells_fast(table)
        if not all_rows:
            return table_cars

        # 获取并处理表头
        headers = [clean_text(cell) for cell in all_rows[0] if cell]
        if not headers:
            return table_cars

        # 显示表格结构信息（仅在详细模式下）
        if self.verbose:
            console.print(f"\n[cyan]表格 {table_index + 1} 结构信息:[/cyan]")
            console.print(f"表头: {headers}")
            console.print(f"总行数: {len(all_rows)}")
            if len(all_rows) > 1:
                console.print(f"第一行数据示例: {all_rows[1]}")

        # 根据表头判断表格类型
        current_subsection_title = (
            self.current_subsection.title if self.current_subsection else None
        )
        self.logger.debug(
            f"表格 {table_index + 1} 上下文信息: 分类={self.current_category}, 子分类={current_subsection_title}"
        )

        table_category, table_type = get_table_type(
            headers, self.current_category, current_subsection_title
        )

        # 如果当前在子分类中，确保使用正确的子分类名称
        if self.current_subsection and table_type == "未知":
            table_type = self.current_subsection.title
            self.logger.debug(f"表格 {table_index + 1} 使用当前子分类: {table_type}")

        self.logger.info(
            f"表格 {table_index + 1} 类型判断: 分类={table_category}, 子分类={table_type}"
        )

        # 预先创建基础信息
        base_info = {
            "category": table_category,
            "sub_type": table_type,
            "car_type": 2 if table_category == "节能型" else 1,
            "batch": batch_number,
            "table_id": table_index + 1,  # 添加表格ID, 从1开始计数
        }

        total_rows = len(all_rows) - 1
        if total_rows > 100:
            console.print(f"[dim]开始处理大表格, 共 {total_rows} 行[/dim]")

        # 分块处理数据行
        for chunk_start in range(1, len(all_rows), self._chunk_size):
            chunk_end = min(chunk_start + self._chunk_size, len(all_rows))
            chunk_rows = all_rows[chunk_start:chunk_end]

            # 批量处理当前块的数据行
            for row_idx, cells in enumerate(chunk_rows, chunk_start):
                # 跳过空行
                if not any(str(cell).strip() for cell in cells):
                    continue

                # 记录列数不匹配的情况, 但仍然处理数据
                if len(cells) != len(headers):
                    if self.verbose:
                        console.print(
                            f"[yellow]表格 {table_index + 1} 第 {row_idx} 行列数不匹配: "
                            f"预期 {len(headers)} 列, 实际 {len(cells)} 列[/yellow]"
                        )
                        console.print(f"行内容: {cells}")
                    # 调整单元格数量以匹配表头
                    if len(cells) > len(headers):
                        cells = cells[: len(headers)]
                    else:
                        cells.extend([""] * (len(headers) - len(cells)))

                # 创建新的字典, 避免引用同一个对象
                car_info = base_info.copy()
                car_info["raw_text"] = " | ".join(str(cell) for cell in cells)

                # 使用zip优化字段映射, 同时清理文本
                car_info.update(
                    {
                        header: clean_text(str(value))
                        for header, value in zip(headers, cells)
                    }
                )

                # 处理车辆信息
                car_info = process_car_info(car_info, batch_number)
                table_cars.append(car_info)

            if total_rows > 100:
                progress = (chunk_end - 1) / total_rows * 100
                console.print(
                    f"[dim]处理进度: {progress:.1f}% ({chunk_end-1}/{total_rows})[/dim]"
                )

            # 主动触发垃圾回收
            if len(table_cars) > 5000:
                gc.collect()

        # 缓存结果
        self._table_cache[table_index] = table_cars

        # 记录处理时间和统计信息
        elapsed = time.time() - start_time
        if total_rows > 100 or len(table_cars) > 0:
            console.print(
                f"[dim]表格 {table_index + 1} 处理了 {total_rows} 行, "
                f"数据 {len(table_cars)} 行, 耗时: {elapsed:.2f}秒[/dim]"
            )

        return table_cars

    def _log_time(self, operation: str) -> None:
        """记录操作耗时"""
        current_time = time.time()
        elapsed = current_time - self.start_time
        self._processing_times[operation] = elapsed
        if operation != "init" and self.verbose:
            console.print(f"[dim]{operation} 耗时: {elapsed:.2f}秒[/dim]")
        self.start_time = current_time

    # @profile_function
    def process(self) -> List[Dict[str, Any]]:
        """处理文档并返回所有车辆信息"""
        try:
            self.logger.info(f"开始处理文档: {self.doc_path}")
            self._log_time("init")

            table_count = 0
            row_count = 0
            error_count = 0

            # 遍历文档中的所有元素
            for element in self.doc.element.body:
                try:
                    # 处理段落
                    if element.tag.endswith("p"):
                        text = element.text.strip()
                        if not text:
                            continue

                        # 提取批次号
                        if not self.batch_number:
                            self.batch_number = extract_batch_number(text)
                            if self.batch_number:
                                self.doc_structure.set_batch_number(self.batch_number)
                                self.logger.info(f"提取到批次号: {self.batch_number}")
                                self.doc_structure.add_node(
                                    f"第{self.batch_number}批", "batch", level=0
                                )

                        # 更新分类信息
                        if "节能型汽车" in text:
                            self.current_category = "节能型"
                            self.current_section = self.doc_structure.add_node(
                                "节能型汽车", "section", content=text
                            )
                            self.current_subsection = None
                            self.current_numbered_section = None
                            self.logger.debug(f"更新分类: {self.current_category}")
                        elif "新能源汽车" in text:
                            self.current_category = "新能源"
                            self.current_section = self.doc_structure.add_node(
                                "新能源汽车", "section", content=text
                            )
                            self.current_subsection = None
                            self.current_numbered_section = None
                            self.logger.debug(f"更新分类: {self.current_category}")
                        elif text.startswith("（") and not any(
                            str.isdigit() for str in text
                        ):
                            self.current_subsection = self.doc_structure.add_node(
                                text.strip(),
                                "subsection",
                                content=text,
                                parent_node=self.current_section,
                            )
                            self.current_numbered_section = None
                            self.logger.debug(f"更新类型: {text}")
                        # 处理带数字编号的节点
                        elif text.startswith(("1.", "2.", "3.", "4.", "5.")):
                            self.current_numbered_section = self.doc_structure.add_node(
                                text.strip(),
                                "numbered_section",
                                content=text,
                                parent_node=self.current_subsection
                                or self.current_section,
                            )
                            self.logger.debug(f"更新编号节点: {text}")
                        # 处理带括号数字编号的子节点
                        elif text.startswith("（") and any(
                            num in text for num in "123456789"
                        ):
                            if self.current_numbered_section:
                                self.doc_structure.add_node(
                                    text.strip(),
                                    "numbered_subsection",
                                    content=text,
                                    parent_node=self.current_numbered_section,
                                )
                            else:
                                self.doc_structure.add_node(
                                    text.strip(),
                                    "numbered_subsection",
                                    content=text,
                                    parent_node=self.current_subsection
                                    or self.current_section,
                                )
                            self.logger.debug(f"更新编号子节点: {text}")
                        elif "勘误" in text or "说明" in text:
                            self.doc_structure.add_node(
                                text[:20] + "...",
                                "note",
                                content=text,
                                parent_node=self.current_section,
                            )
                        elif "更正" in text or "修改" in text:
                            self.doc_structure.add_node(
                                text[:20] + "...",
                                "correction",
                                content=text,
                                parent_node=self.current_section,
                            )
                        else:
                            self.doc_structure.add_node(
                                text[:20] + "...",
                                "text",
                                content=text,
                                parent_node=self.current_section,
                            )

                    # 处理表格
                    elif element.tag.endswith("tbl"):
                        table_count += 1
                        for i, table in enumerate(self.doc.tables):
                            if table._element is element:
                                if table.rows:
                                    row_count += len(table.rows)
                                try:
                                    table_cars = self._extract_car_info(
                                        i, self.batch_number
                                    )
                                    self.cars.extend(table_cars)

                                    # 添加表格节点到正确的父节点
                                    parent_node = (
                                        self.current_numbered_section
                                        or self.current_subsection
                                        or self.current_section
                                    )
                                    self.doc_structure.add_node(
                                        f"表格 {i+1}",
                                        "table",
                                        metadata={
                                            "rows": len(table.rows),
                                            "columns": len(table.rows[0].cells)
                                            if table.rows
                                            else 0,
                                            "records": len(table_cars),
                                            "category": self.current_category,
                                            "sub_type": self.current_subsection.title
                                            if self.current_subsection
                                            else None,
                                        },
                                        parent_node=parent_node,
                                    )

                                    if self.verbose:
                                        self.logger.info(
                                            f"处理表格 {i+1}, 提取到 {len(table_cars)} 条记录"
                                        )
                                except Exception as e:
                                    error_count += 1
                                    self.logger.error(f"处理表格 {i+1} 出错: {str(e)}")
                                break
                except Exception as e:
                    error_count += 1
                    self.logger.error(f"处理元素出错: {str(e)}")
                    continue

            self._log_time("process")
            self.logger.info(
                f"文档处理完成: {table_count} 个表格, {row_count} 行, "
                f"{len(self.cars)} 条记录, {error_count} 个错误"
            )

            # 执行批次数据一致性验证 - 只在处理后执行一次
            verification_start = time.time()
            consistency_result = self.verify_batch_consistency()
            verification_time = time.time() - verification_start
            self.logger.info(
                f"批次一致性验证结果: {consistency_result['status']} (耗时: {verification_time:.2f}秒)"
            )

            # 计算文件大小和记录数
            file_size = os.path.getsize(self.doc_path) / (1024 * 1024)  # MB
            record_count = len(self.cars)

            # 对于大文件或大量记录, 禁用详细显示以提高性能
            is_large_file = file_size > 50 or record_count > 100000  # 50MB或1万条记录

            # 显示文档结构（仅在详细模式下）
            if self.verbose and not is_large_file:
                display_doc_content(self.doc_structure)
            elif self.verbose and is_large_file:
                console.print(
                    "[yellow]文件较大, 跳过显示详细文档结构以提高性能[/yellow]"
                )

            # 显示批次一致性验证结果（始终显示, 即使在简洁模式下）
            self._display_consistency_result(consistency_result)

            # 处理完成后主动释放资源
            self._table_cache.clear()
            gc.collect()

            return self.cars

        except Exception as e:
            self.logger.error(f"处理文档失败: {str(e)}")
            raise ProcessingError(f"处理文档 {self.doc_path} 失败: {str(e)}")

    def verify_batch_consistency(self) -> dict:
        """
        验证每个批次的表格数据总和是否与批次总记录数一致
        即便没有声明的总记录数, 也验证表格记录数与处理后的记录数是否一致

        Returns:
            dict: 包含批次验证结果的字典
        """
        # 如果没有批次号, 直接返回
        if not self.batch_number:
            return {"status": "no_batch", "message": "未找到批次号"}

        # 按表格分组统计车辆记录数
        table_counts = {}
        for car in self.cars:
            table_id = car.get("table_id", "未知")
            if table_id not in table_counts:
                table_counts[table_id] = 0
            table_counts[table_id] += 1

        # 计算从表格中提取的总记录数
        total_extracted_count = sum(table_counts.values())

        # 获取批次声明的总记录数
        if self.declared_count is None:
            self.declared_count = self._extract_declared_count()

        # 验证结果
        if self.declared_count is not None:
            # 如果有声明的总记录数, 比较声明数与实际数
            if total_extracted_count == self.declared_count:
                return {
                    "status": "match",
                    "message": f"批次记录数匹配：声明 {self.declared_count}, 实际 {total_extracted_count}",
                    "batch": self.batch_number,
                    "actual_count": total_extracted_count,
                    "declared_count": self.declared_count,
                    "table_counts": table_counts,
                }
            else:
                return {
                    "status": "mismatch",
                    "message": f"批次记录数不匹配：声明 {self.declared_count}, 实际 {total_extracted_count}",
                    "batch": self.batch_number,
                    "actual_count": total_extracted_count,
                    "declared_count": self.declared_count,
                    "table_counts": table_counts,
                    "difference": self.declared_count - total_extracted_count,
                }
        else:
            # 如果没有声明的总记录数, 验证表格总记录数与处理后的记录数是否一致
            processed_count = len(self.cars)
            if total_extracted_count == processed_count:
                return {
                    "status": "internal_match",
                    "message": f"内部一致性检查通过：表格记录总数 {total_extracted_count} 与处理结果数 {processed_count} 一致",
                    "batch": self.batch_number,
                    "actual_count": total_extracted_count,
                    "processed_count": processed_count,
                    "table_counts": table_counts,
                }
            else:
                return {
                    "status": "internal_mismatch",
                    "message": f"内部一致性检查失败：表格记录总数 {total_extracted_count} 与处理结果数 {processed_count} 不一致",
                    "batch": self.batch_number,
                    "actual_count": total_extracted_count,
                    "processed_count": processed_count,
                    "table_counts": table_counts,
                    "difference": total_extracted_count - processed_count,
                }

    def _display_consistency_result(self, result: Dict[str, Any]) -> None:
        """显示批次一致性验证结果"""
        # 在显示结果前添加标题, 表明这是关键信息
        console.print()
        console.print("[bold cyan]📊 关键信息：数据一致性检查[/bold cyan]")

        if result["status"] == "no_batch":
            console.print(
                Panel(
                    "[yellow]⚠️ 未找到批次号, 无法验证数据一致性[/yellow]",
                    title="数据一致性检查",
                    border_style="yellow",
                )
            )
            return

        if result["status"] == "unknown":
            console.print(
                Panel(
                    f"[yellow]⚠️ 第{result['batch']}批：未找到总记录数声明, 实际记录数为 {result['actual_count']}[/yellow]",
                    title="数据一致性检查",
                    border_style="yellow",
                )
            )
        elif result["status"] == "match":
            console.print(
                Panel(
                    f"[green]✅ 第{result['batch']}批：记录数匹配, 共 {result['actual_count']} 条记录[/green]",
                    title="数据一致性检查",
                    border_style="green",
                )
            )
        elif result["status"] == "mismatch":
            diff_text = (
                f"差异 {result['difference']} 条" if "difference" in result else ""
            )
            console.print(
                Panel(
                    f"[red]❌ 第{result['batch']}批：记录数不匹配！声明 {result['declared_count']}, 实际 {result['actual_count']}, {diff_text}[/red]",
                    title="⚠️ 数据一致性检查",
                    border_style="red",
                )
            )
        elif result["status"] == "internal_match":
            console.print(
                Panel(
                    f"[green]✅ 第{result['batch']}批：内部一致性检查通过, 表格记录总数 {result['actual_count']} 与处理结果数 {result['processed_count']} 一致[/green]",
                    title="数据一致性检查",
                    border_style="green",
                )
            )
        elif result["status"] == "internal_mismatch":
            diff_text = (
                f"差异 {result['difference']} 条" if "difference" in result else ""
            )
            console.print(
                Panel(
                    f"[red]❌ 第{result['batch']}批：内部一致性检查失败！表格记录总数 {result['actual_count']} 与处理结果数 {result['processed_count']} 不一致, {diff_text}[/red]",
                    title="⚠️ 数据一致性检查",
                    border_style="red",
                )
            )

        # 显示表格记录分布
        table_counts = result.get("table_counts", {})
        if table_counts:
            count_table = Table(
                title="📊 表格记录分布",
                title_style="bold cyan",
                show_header=True,
                header_style="bold green",
                border_style="blue",
            )
            count_table.add_column("表格ID", style="cyan")
            count_table.add_column("记录数", justify="right", style="green")
            count_table.add_column("占比", justify="right", style="yellow")

            total = result.get("actual_count", sum(table_counts.values()))

            for table_id, count in sorted(table_counts.items()):
                percentage = (count / total * 100) if total > 0 else 0
                count_table.add_row(
                    f"表格 {table_id}" if not isinstance(table_id, str) else table_id,
                    str(count),
                    f"{percentage:.1f}%",
                )

            console.print(count_table)


@cli.command()
@click.argument(
    "input_path",
    type=click.Path(exists=True),
)
@click.option(
    "-o",
    "--output",
    type=click.Path(dir_okay=False),
    default="cars_output.csv",
    help="输出CSV文件路径",
)
@click.option("-v", "--verbose", is_flag=True, help="显示详细处理信息")
@click.option("--preview", is_flag=True, help="显示文档内容预览")
@click.option(
    "--compare",
    type=click.Path(exists=True, dir_okay=False),
    help="与指定的CSV文件进行对比",
)
@click.option(
    "--config",
    type=click.Path(exists=True, dir_okay=False),
    help="配置文件路径",
)
def process(
    input_path: str,
    output: str,
    verbose: bool,
    preview: bool,
    compare: Optional[str] = None,
    config: Optional[str] = None,
) -> None:
    """处理指定的docx文件或目录下的所有docx文件"""
    try:
        # 设置日志
        setup_logging()
        logger = logging.getLogger(__name__)
        logger.info(f"开始处理任务: 输入={input_path}, 输出={output}")

        # 加载配置
        config_data = {}
        if config:
            try:
                config_data = load_config(config)
                logger.info(f"加载配置文件: {config}")
            except ConfigurationError as e:
                logger.error(f"加载配置失败: {str(e)}")
                console.print(f"[bold red]加载配置失败: {str(e)}")
                return

        process_files(input_path, output, verbose, preview, compare, config_data)

    except Exception as e:
        logger.error(f"处理任务失败: {str(e)}")
        console.print(f"[bold red]处理任务失败: {str(e)}")


if __name__ == "__main__":
    cli()
